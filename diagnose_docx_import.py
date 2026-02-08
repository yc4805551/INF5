#!/usr/bin/env python3
"""
DOCX å¯¼å…¥åŠŸèƒ½ - ä¸€é”®éƒ¨ç½²å’Œè¯Šæ–­è„šæœ¬
è¿è¡Œæ­¤è„šæœ¬å°†è‡ªåŠ¨å®Œæˆéƒ¨ç½²å¹¶æµ‹è¯• DOCX å¯¼å…¥åŠŸèƒ½
"""

import os
import sys
import subprocess
import requests
import json
from datetime import datetime
from pathlib import Path

# é…ç½®
BACKEND_URL = "http://localhost:5000"
PROJECT_ROOT = Path(__file__).parent.absolute()

class Logger:
    """å½©è‰²æ—¥å¿—è¾“å‡º"""
    COLORS = {
        'HEADER': '\033[95m',
        'BLUE': '\033[94m',
        'GREEN': '\033[92m',
        'YELLOW': '\033[93m',
        'RED': '\033[91m',
        'END': '\033[0m',
        'BOLD': '\033[1m',
    }
    
    @staticmethod
    def header(msg):
        print(f"\n{Logger.COLORS['HEADER']}{Logger.COLORS['BOLD']}{'='*60}{Logger.COLORS['END']}")
        print(f"{Logger.COLORS['HEADER']}{Logger.COLORS['BOLD']}{msg}{Logger.COLORS['END']}")
        print(f"{Logger.COLORS['HEADER']}{Logger.COLORS['BOLD']}{'='*60}{Logger.COLORS['END']}\n")
    
    @staticmethod
    def info(msg):
        print(f"{Logger.COLORS['BLUE']}â„¹ï¸  {msg}{Logger.COLORS['END']}")
    
    @staticmethod
    def success(msg):
        print(f"{Logger.COLORS['GREEN']}âœ… {msg}{Logger.COLORS['END']}")
    
    @staticmethod
    def warning(msg):
        print(f"{Logger.COLORS['YELLOW']}âš ï¸  {msg}{Logger.COLORS['END']}")
    
    @staticmethod
    def error(msg):
        print(f"{Logger.COLORS['RED']}âŒ {msg}{Logger.COLORS['END']}")
    
    @staticmethod
    def step(num, msg):
        print(f"\n{Logger.COLORS['BOLD']}æ­¥éª¤ {num}: {msg}{Logger.COLORS['END']}")

def run_command(cmd, cwd=None, check=True):
    """æ‰§è¡Œç³»ç»Ÿå‘½ä»¤å¹¶è¿”å›è¾“å‡º"""
    try:
        result = subprocess.run(
            cmd,
            cwd=cwd or PROJECT_ROOT,
            shell=True,
            capture_output=True,
            text=True,
            timeout=30
        )
        if check and result.returncode != 0:
            Logger.error(f"å‘½ä»¤æ‰§è¡Œå¤±è´¥: {cmd}")
            Logger.error(f"é”™è¯¯è¾“å‡º: {result.stderr}")
            return None
        return result
    except subprocess.TimeoutExpired:
        Logger.error(f"å‘½ä»¤è¶…æ—¶: {cmd}")
        return None
    except Exception as e:
        Logger.error(f"å‘½ä»¤æ‰§è¡Œå¼‚å¸¸: {e}")
        return None

def check_git_status():
    """æ£€æŸ¥ Git çŠ¶æ€"""
    Logger.step(1, "æ£€æŸ¥ Git ä»“åº“çŠ¶æ€")
    
    # æ£€æŸ¥æ˜¯å¦åœ¨æ­£ç¡®çš„ç›®å½•
    if not (PROJECT_ROOT / '.git').exists():
        Logger.error(f"å½“å‰ç›®å½•ä¸æ˜¯ Git ä»“åº“: {PROJECT_ROOT}")
        return False
    
    # è·å–å½“å‰åˆ†æ”¯
    result = run_command("git branch --show-current")
    if result:
        branch = result.stdout.strip()
        Logger.info(f"å½“å‰åˆ†æ”¯: {branch}")
    
    # è·å–æœ€æ–°æäº¤
    result = run_command("git log -1 --oneline")
    if result:
        Logger.info(f"å½“å‰æäº¤: {result.stdout.strip()}")
    
    return True

def deploy_latest_code():
    """æ‹‰å–æœ€æ–°ä»£ç """
    Logger.step(2, "æ‹‰å–æœ€æ–°ä»£ç ")
    
    # è·å–è¿œç¨‹æ›´æ–°
    Logger.info("æ­£åœ¨è·å–è¿œç¨‹æ›´æ–°...")
    result = run_command("git fetch origin")
    if not result:
        return False
    
    # æ£€æŸ¥æ˜¯å¦æœ‰æœªæäº¤çš„æ›´æ”¹
    result = run_command("git status --porcelain")
    if result and result.stdout.strip():
        Logger.warning("æ£€æµ‹åˆ°æœªæäº¤çš„æœ¬åœ°æ›´æ”¹")
        print(result.stdout)
        response = input("æ˜¯å¦è¦æš‚å­˜è¿™äº›æ›´æ”¹å¹¶ç»§ç»­? (y/n): ")
        if response.lower() == 'y':
            run_command("git stash")
            Logger.info("å·²æš‚å­˜æœ¬åœ°æ›´æ”¹")
        else:
            Logger.error("éƒ¨ç½²å·²å–æ¶ˆ")
            return False
    
    # æ‹‰å–æœ€æ–°ä»£ç 
    Logger.info("æ­£åœ¨æ‹‰å–æœ€æ–°ä»£ç ...")
    result = run_command("git pull origin main")
    if not result:
        return False
    
    if "Already up to date" in result.stdout:
        Logger.success("ä»£ç å·²æ˜¯æœ€æ–°ç‰ˆæœ¬")
    else:
        Logger.success("ä»£ç æ›´æ–°æˆåŠŸ")
        print(result.stdout)
    
    # æ˜¾ç¤ºæœ€æ–°çš„æäº¤
    result = run_command("git log -1 --oneline")
    if result:
        Logger.info(f"æœ€æ–°æäº¤: {result.stdout.strip()}")
    
    return True

def check_backend_service():
    """æ£€æŸ¥åç«¯æœåŠ¡çŠ¶æ€"""
    Logger.step(3, "æ£€æŸ¥åç«¯æœåŠ¡")
    
    try:
        response = requests.get(f"{BACKEND_URL}/health", timeout=5)
        if response.status_code == 200:
            Logger.success(f"åç«¯æœåŠ¡è¿è¡Œæ­£å¸¸ ({BACKEND_URL})")
            return True
    except requests.exceptions.ConnectionError:
        Logger.error("åç«¯æœåŠ¡æœªå“åº”")
    except requests.exceptions.Timeout:
        Logger.error("åç«¯æœåŠ¡å“åº”è¶…æ—¶")
    except Exception as e:
        Logger.error(f"æ£€æŸ¥åç«¯æœåŠ¡æ—¶å‡ºé”™: {e}")
    
    Logger.warning("è¯·æ‰‹åŠ¨é‡å¯åç«¯æœåŠ¡:")
    Logger.info("  æ–¹æ³•1: cd backend && python app.py")
    Logger.info("  æ–¹æ³•2: systemctl restart infv5-backend")
    
    response = input("\næ˜¯å¦å·²é‡å¯åç«¯æœåŠ¡? (y/n): ")
    return response.lower() == 'y'

def test_docx_import():
    """æµ‹è¯• DOCX å¯¼å…¥åŠŸèƒ½"""
    Logger.step(4, "æµ‹è¯• DOCX å¯¼å…¥åŠŸèƒ½")
    
    # åˆ›å»ºæµ‹è¯• DOCX æ–‡ä»¶
    Logger.info("åˆ›å»ºæµ‹è¯• DOCX æ–‡ä»¶...")
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('æµ‹è¯•æ–‡æ¡£', 0)
        doc.add_paragraph('è¿™æ˜¯ç¬¬ä¸€æ®µæµ‹è¯•æ–‡æœ¬ã€‚')
        doc.add_paragraph('è¿™æ˜¯ç¬¬äºŒæ®µæµ‹è¯•æ–‡æœ¬ï¼ŒåŒ…å«ä¸€äº›**ç²—ä½“**å†…å®¹ã€‚')
        
        # æ·»åŠ è¡¨æ ¼
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = 'åˆ—1'
        table.rows[0].cells[1].text = 'åˆ—2'
        table.rows[1].cells[0].text = 'æ•°æ®1'
        table.rows[1].cells[1].text = 'æ•°æ®2'
        
        test_file = PROJECT_ROOT / 'test_docx_import.docx'
        doc.save(test_file)
        Logger.success(f"æµ‹è¯•æ–‡ä»¶å·²åˆ›å»º: {test_file}")
    except ImportError:
        Logger.error("éœ€è¦å®‰è£… python-docx: pip install python-docx")
        return False
    except Exception as e:
        Logger.error(f"åˆ›å»ºæµ‹è¯•æ–‡ä»¶å¤±è´¥: {e}")
        return False
    
    # æµ‹è¯•å¯¼å…¥ API
    Logger.info("è°ƒç”¨åç«¯å¯¼å…¥ API...")
    try:
        with open(test_file, 'rb') as f:
            files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = requests.post(
                f"{BACKEND_URL}/api/canvas/import-from-docx",
                files=files,
                timeout=10
            )
        
        Logger.info(f"HTTP çŠ¶æ€ç : {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            Logger.success("å¯¼å…¥æˆåŠŸ!")
            
            # åˆ†æè¿”å›æ•°æ®
            content = data.get('content', [])
            Logger.info(f"è§£æå‡º {len(content)} ä¸ªå†…å®¹èŠ‚ç‚¹")
            
            # æ£€æŸ¥ç©ºæ–‡æœ¬èŠ‚ç‚¹
            empty_nodes = []
            def check_nodes(nodes, path=""):
                for i, node in enumerate(nodes):
                    node_path = f"{path}[{i}]"
                    if node.get('type') == 'text' and not node.get('text'):
                        empty_nodes.append(node_path)
                    if 'content' in node:
                        check_nodes(node['content'], f"{node_path}.content")
            
            check_nodes(content)
            
            if empty_nodes:
                Logger.warning(f"å‘ç° {len(empty_nodes)} ä¸ªç©ºæ–‡æœ¬èŠ‚ç‚¹ï¼ˆåº”ç”±å‰ç«¯æ¸…æ´—ï¼‰:")
                for path in empty_nodes[:5]:  # åªæ˜¾ç¤ºå‰5ä¸ª
                    Logger.warning(f"  - {path}")
            else:
                Logger.success("æœªå‘ç°ç©ºæ–‡æœ¬èŠ‚ç‚¹ - æ•°æ®è´¨é‡è‰¯å¥½")
            
            # æ˜¾ç¤ºç¤ºä¾‹å†…å®¹
            Logger.info("\nè¿”å›æ•°æ®ç¤ºä¾‹ (å‰3ä¸ªèŠ‚ç‚¹):")
            print(json.dumps(content[:3], indent=2, ensure_ascii=False))
            
            return True
        else:
            Logger.error(f"å¯¼å…¥å¤±è´¥: {response.status_code}")
            Logger.error(f"é”™è¯¯ä¿¡æ¯: {response.text}")
            return False
            
    except requests.exceptions.ConnectionError:
        Logger.error("æ— æ³•è¿æ¥åˆ°åç«¯æœåŠ¡")
        return False
    except Exception as e:
        Logger.error(f"æµ‹è¯•å¯¼å…¥æ—¶å‡ºé”™: {e}")
        return False
    finally:
        # æ¸…ç†æµ‹è¯•æ–‡ä»¶
        if test_file.exists():
            test_file.unlink()

def check_frontend_files():
    """æ£€æŸ¥å…³é”®å‰ç«¯æ–‡ä»¶æ˜¯å¦å­˜åœ¨ä¿®å¤"""
    Logger.step(5, "éªŒè¯å‰ç«¯ä¿®å¤æ–‡ä»¶")
    
    checks = {
        'frontend/src/features/fast-canvas/FastCanvasView.tsx': 'sanitizeContent',
        'frontend/src/features/fast-canvas/components/Editor/TiptapEditor.css': '--tw-prose-body',
        'backend/features/canvas_converter.py': 'fallback'
    }
    
    all_ok = True
    for file_path, keyword in checks.items():
        full_path = PROJECT_ROOT / file_path
        if not full_path.exists():
            Logger.error(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")
            all_ok = False
            continue
        
        with open(full_path, 'r', encoding='utf-8') as f:
            content = f.read()
            if keyword in content:
                Logger.success(f"âœ“ {file_path} åŒ…å«ä¿®å¤: {keyword}")
            else:
                Logger.warning(f"âœ— {file_path} å¯èƒ½æœªåŒ…å«ä¿®å¤: {keyword}")
                all_ok = False
    
    return all_ok

def generate_summary(results):
    """ç”Ÿæˆè¯Šæ–­æ‘˜è¦"""
    Logger.header("è¯Šæ–­æ‘˜è¦")
    
    print(f"æ—¶é—´: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"é¡¹ç›®è·¯å¾„: {PROJECT_ROOT}")
    print(f"åç«¯åœ°å€: {BACKEND_URL}")
    print()
    
    status_map = {
        'git': results.get('git', False),
        'deploy': results.get('deploy', False),
        'backend': results.get('backend', False),
        'frontend': results.get('frontend', False),
        'import': results.get('import', False),
    }
    
    for key, status in status_map.items():
        icon = "âœ…" if status else "âŒ"
        print(f"{icon} {key.upper()}: {'é€šè¿‡' if status else 'å¤±è´¥'}")
    
    all_passed = all(status_map.values())
    
    if all_passed:
        Logger.header("ğŸ‰ æ‰€æœ‰æ£€æŸ¥é€šè¿‡ï¼DOCX å¯¼å…¥åŠŸèƒ½æ­£å¸¸")
        Logger.info("æ‚¨ç°åœ¨å¯ä»¥åœ¨æµè§ˆå™¨ä¸­ä½¿ç”¨ DOCX å¯¼å…¥åŠŸèƒ½äº†")
        Logger.info("å¦‚ä»æœ‰é—®é¢˜ï¼Œè¯·æ¸…é™¤æµè§ˆå™¨ç¼“å­˜ (Ctrl+F5)")
    else:
        Logger.header("âš ï¸  éƒ¨åˆ†æ£€æŸ¥æœªé€šè¿‡")
        Logger.info("è¯·å°†æ­¤æ—¥å¿—è¾“å‡ºå‘é€ç»™å¼€å‘è€…ä»¥è·å¾—å¸®åŠ©")

def main():
    """ä¸»å‡½æ•°"""
    Logger.header("DOCX å¯¼å…¥åŠŸèƒ½ - è‡ªåŠ¨éƒ¨ç½²å’Œè¯Šæ–­")
    
    results = {}
    
    # æ£€æŸ¥ Git çŠ¶æ€
    results['git'] = check_git_status()
    if not results['git']:
        Logger.error("Git æ£€æŸ¥å¤±è´¥ï¼Œè¯·ç¡®ä¿åœ¨é¡¹ç›®æ ¹ç›®å½•è¿è¡Œæ­¤è„šæœ¬")
        return 1
    
    # éƒ¨ç½²æœ€æ–°ä»£ç 
    results['deploy'] = deploy_latest_code()
    if not results['deploy']:
        Logger.error("ä»£ç éƒ¨ç½²å¤±è´¥")
        return 1
    
    # æ£€æŸ¥åç«¯æœåŠ¡
    results['backend'] = check_backend_service()
    if not results['backend']:
        Logger.error("åç«¯æœåŠ¡æ£€æŸ¥å¤±è´¥")
        return 1
    
    # éªŒè¯å‰ç«¯æ–‡ä»¶
    results['frontend'] = check_frontend_files()
    
    # æµ‹è¯• DOCX å¯¼å…¥
    results['import'] = test_docx_import()
    
    # ç”Ÿæˆæ‘˜è¦
    generate_summary(results)
    
    return 0 if all(results.values()) else 1

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        Logger.warning("\nç”¨æˆ·å–æ¶ˆæ“ä½œ")
        sys.exit(1)
    except Exception as e:
        Logger.error(f"\næœªé¢„æœŸçš„é”™è¯¯: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
