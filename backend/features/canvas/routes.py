import re
from flask import Blueprint, request, jsonify, Response, send_file, current_app
import logging
import io
import os
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from core.services import current_engine, llm_engine
from core.win32_engine import WordAppEngine

canvas_bp = Blueprint('canvas', __name__)

@canvas_bp.route('/upload', methods=['POST'])
def canvas_upload():
    try:
        logging.info("Starting canvas_upload...")
        if 'file' not in request.files:
            logging.error("Canvas Upload: No file part")
            return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        if file.filename == '':
            logging.error("Canvas Upload: No selected file")
            return jsonify({"error": "No selected file"}), 400
        
        logging.info(f"Canvas Upload: processing file {file.filename}")
        content = file.read()
        logging.info(f"Canvas Upload: file size {len(content)} bytes")
        
        current_engine.load_document(io.BytesIO(content))
        logging.info("Canvas Upload: load_document successful")
        
        # Pagination defaults
        page_size = 100
        total_paras = current_engine.get_paragraph_count()
        
        preview = current_engine.get_preview_data(start=0, limit=page_size)
        logging.info(f"Canvas Upload: preview items count: {len(preview)}")
        
        html_preview = current_engine.get_html_preview(limit=page_size)
        total_paragraphs = current_engine.get_paragraph_count()
        structure = current_engine.get_document_structure()
        
        return jsonify({
            "message": "File uploaded successfully", 
            "filename": file.filename,
            "html_preview": html_preview,
            "total_paragraphs": total_paragraphs,
            "page_size": page_size,
            "structure": structure
        }), 200
    except Exception as e:
        import traceback
        logging.error(f"Canvas Upload Error: {e}")
        logging.error(traceback.format_exc())
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/preview', methods=['GET'])
def canvas_preview():
    try:
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('page_size', 100))
        start = (page - 1) * page_size
        return jsonify(current_engine.get_preview_data(start=start, limit=page_size))
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/preview_html', methods=['GET'])
def canvas_preview_html():
    try:
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('page_size', 100))
        start = (page - 1) * page_size
        return jsonify({
            "html": current_engine.get_html_preview(start=start, limit=page_size),
            "total_paragraphs": current_engine.get_paragraph_count(),
            "page": page,
            "page_size": page_size,
            "structure": current_engine.get_document_structure()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/upload_reference', methods=['POST'])
def canvas_upload_reference():
    try:
        logging.info("Starting canvas_upload_reference...")
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        content = file.read()
        success, message = current_engine.add_reference_doc(io.BytesIO(content), file.filename)
        
        if success:
            refs = current_engine.get_reference_list()
            return jsonify({"message": f"Reference {file.filename} added", "references": refs})
        else:
            return jsonify({"error": f"Failed to load reference: {message}"}), 500
    except Exception as e:
        logging.error(f"Canvas Upload Reference Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/remove_reference', methods=['POST'])
def canvas_remove_reference():
    try:
        data = request.get_json()
        filename = data.get("filename")
        if not filename:
            return jsonify({"error": "Filename required"}), 400
            
        success = current_engine.remove_reference_doc(filename)
        refs = current_engine.get_reference_list()
        
        if success:
            return jsonify({"message": f"Removed {filename}", "references": refs})
        else:
            return jsonify({"error": "File not found", "references": refs}), 404
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/references', methods=['GET'])
def canvas_get_references():
    try:
        refs = current_engine.get_reference_list()
        return jsonify({"references": refs})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/reset', methods=['POST'])
def canvas_reset():
    try:
        current_engine.reset()
        return jsonify({"message": "Canvas reset successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/chat', methods=['POST'])
def canvas_chat():
    try:
        data = request.get_json()
        user_text = data.get("message")
        model_config = data.get("model_config", {})
        history = data.get("history", [])
        selection_context = data.get("selection_context", [])

        if not user_text:
            return jsonify({"error": "Message required"}), 400

        # Create staging copy if needed
        if not current_engine.staging_doc:
            current_engine.create_staging_copy()

        page = data.get("page", 1)
        page_size = data.get("page_size", 100)
        scope_range = data.get("scope_range") # [start, end]
        
        # Get context (Paginated or Scoped)
        if scope_range and len(scope_range) == 2:
            start_id, end_id = scope_range
            logging.info(f"Using Scope Range: {start_id} - {end_id}")
            # limit is relative to start, so calculate diff
            # Actually get_preview_data slice logic: paragraphs[start : start + limit]
            # so limit = end - start + 1
            limit = end_id - start_id + 1
            context = current_engine.get_preview_data(start=start_id, limit=limit)
        else:
            # Fallback to current page
            start = (page - 1) * page_size
            context = current_engine.get_preview_data(start=start, limit=page_size)
        
        # --- UNIVERSAL ADAPTIVE STRATEGY (MAIN DOC) ---
        total_paras = current_engine.get_paragraph_count()
        main_doc_context = ""
        is_short_doc = total_paras < 300
        
        if is_short_doc:
            # STRATEGY A: Small Doc -> Full Context Injection
            logging.info(f"Adaptive Strategy: Small Doc ({total_paras} paras). Injecting FULL text.")
            full_text = current_engine.get_all_content()
            main_doc_context = f"【当前编辑文档（全文）】:\n{full_text}\n"
        else:
            # STRATEGY B: Large Doc -> Coarse-to-Fine Retrieval
            logging.info(f"Adaptive Strategy: Large Doc ({total_paras} paras). Using Retrieval.")
            
            # 1. Get Main Doc Structure
            main_structure = current_engine.get_document_structure()
            
            # 2. Analyze Relevance via LLM (Reuse logic)
            if main_structure:
                relevant_indices = llm_engine.analyze_toc_relevance(user_text, main_structure, model_config)
                if relevant_indices:
                    retrieved_content = current_engine.get_content_by_indices(relevant_indices)
                    if retrieved_content:
                        main_doc_context = f"【文档相关章节（智能检索）】:\n{retrieved_content}\n"

        # Ref Context (Already Adaptive via docx_engine.get_relevant_reference_context upgrades)
        # We still need to call it.
        # But notice: current_engine.get_relevant_reference_context was updated to be adaptive.
        # However, the code below used llm_engine.analyze_toc_relevance FIRST.
        # We need to respect the internal adaptive logic of get_relevant_reference_context if we fallback?
        # Actually, let's simplify Reference Logic to rely on the Engine's new smarts?
        # Or keep the TOC analysis?
        # The prompt said: "modify routes.py ... to apply threshold logic check".
        # Since we updated get_relevant_reference_context to handle small docs, we can rely on it
        # BUT analyze_toc_relevance is powerful.
        # Let's keep analyze_toc_relevance but maybe we assume the Engine handles the choice?
        # For now, let's keep the existing flow but if it returns nothing, we trust the engine.
        
        # Get reference context
        ref_structure = current_engine.get_reference_structure()
        ref_context_str = ""
        
        # Prepare TOC Summary
        ref_toc_summary = ""
        if ref_structure:
            # Re-implement inline for safety
            ref_toc_summary = "\n【参考文档目录结构】:\n"
            current_doc_idx = -1
            for item in ref_structure:
                if item.get('doc_idx') != current_doc_idx:
                    current_doc_idx = item.get('doc_idx')
                    ref_toc_summary += f"[文档: {item.get('filename')}]\n"
                indent = "  " * (item.get('level', 1) - 1)
                ref_toc_summary += f"{indent}- {item.get('title')}\n"

        # Execute Ref Retrieval
        if ref_structure:
            valid_indices = llm_engine.analyze_toc_relevance(user_text, ref_structure, model_config)
            if valid_indices:
                ref_context_str = current_engine.get_content_by_indices(valid_indices)
            else:
                ref_context_str = current_engine.get_relevant_reference_context(user_text)
        else:
            ref_context_str = current_engine.get_relevant_reference_context(user_text)


        # Combine Contexts
        final_ref_context = (ref_toc_summary + "\n" + ref_context_str) if ref_context_str else ref_toc_summary
        if len(final_ref_context) > 25000:
             final_ref_context = final_ref_context[:25000] + "\n...[参考资料截断]..."

        # Get Global Context (Meta-Summary) if not full doc
        global_context = ""
        if not is_short_doc:
             global_context = current_engine.get_global_context()

        # Construct Final User Message
        # Priority: Main Doc Context -> Global Context -> Original User Text
        
        augmented_user_text = ""
        if global_context:
            augmented_user_text += f"【全书脉络】:\n{global_context}\n\n"
        
        if main_doc_context:
            augmented_user_text += main_doc_context + "\n\n"
            
        augmented_user_text += f"【当前可视区域】:\n(见 doc_context)\n\n【用户指令】:\n{user_text}"

        # Call LLM
        response = llm_engine.chat_with_doc(
            user_message=augmented_user_text, 
            doc_context=context, # Visual Page context
            ref_context=final_ref_context, 
            model_config=model_config,
            history=history,
            selection_context=selection_context
        )

        intent = response.get("intent")
        reply = response.get("reply") or ""
        code = response.get("code")
        
        is_staging = False
        
        if intent == "MODIFY":
            if code:
                # Execute Code
                success, error_msg = current_engine.execute_code(code)
                if not success:
                     reply = f"{reply}\n\n(Error executing changes: {error_msg})"
                     intent = "CHAT"
                else:
                     is_staging = True
            else:
                intent = "CHAT"
        
        return jsonify({
            "message": "Processed", 
            "reply": reply,
            "intent": intent,
            "preview": current_engine.get_preview_data(),
            "html_preview": current_engine.get_html_preview(),
            "is_staging": is_staging
        })
    except Exception as e:
        logging.error(f"Canvas Chat Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/confirm', methods=['POST'])
def canvas_confirm():
    try:
        success = current_engine.commit_staging()
        if not success:
            return jsonify({"error": "No pending changes to confirm"}), 400
        return jsonify({
            "message": "Changes confirmed",
            "preview": current_engine.get_preview_data(),
            "html_preview": current_engine.get_html_preview()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/discard', methods=['POST'])
def canvas_discard():
    try:
        current_engine.discard_staging()
        return jsonify({
            "message": "Changes discarded",
            "preview": current_engine.get_preview_data(),
            "html_preview": current_engine.get_html_preview()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/download', methods=['GET'])
def canvas_download():
    try:
        stream = current_engine.save_to_stream()
        return send_file(
            stream, 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="modified.docx"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/export_pdf', methods=['GET'])
def canvas_export_pdf():
    try:
        engine = WordAppEngine()
        # Pre-check connection
        if not engine.connect():
             return jsonify({"error": "Failed to connect to Word. Please ensure Word is installed on the server."}), 500
        
        # 1. Save current doc to a temporary file
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) # backend/
        temp_dir = os.path.join(base_dir, 'static', 'temp')
        if not os.path.exists(temp_dir): os.makedirs(temp_dir, exist_ok=True)
        
        import uuid
        unique_id = uuid.uuid4()
        docx_path = os.path.join(temp_dir, f"temp_{unique_id}.docx")
        pdf_path = os.path.join(temp_dir, f"export_{unique_id}.pdf")
        
        current_engine.save_to_path(docx_path)
        
        # 2. Convert to PDF
        success = engine.export_to_pdf(docx_path, pdf_path)
        engine.quit() # Always cleanup
        
        if success and os.path.exists(pdf_path):
             return send_file(
                pdf_path, 
                mimetype="application/pdf",
                as_attachment=True, 
                download_name="document.pdf"
             )
        else:
             return jsonify({"error": "PDF conversion failed in Word engine."}), 500
             
    except Exception as e:
        logging.error(f"PDF Export Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/update_toc', methods=['POST'])
def canvas_update_toc():
    try:
        engine = WordAppEngine()
        if not engine.connect():
             return jsonify({"error": "Failed to connect to Word."}), 500
             
        # 1. Save current state
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        temp_dir = os.path.join(base_dir, 'static', 'temp')
        if not os.path.exists(temp_dir): os.makedirs(temp_dir, exist_ok=True)
        
        import uuid
        unique_id = uuid.uuid4()
        docx_path = os.path.join(temp_dir, f"temp_toc_{unique_id}.docx")
        
        current_engine.save_to_path(docx_path)
        
        # 2. Update TOC
        success = engine.update_toc(docx_path)
        engine.quit()
        
        if success:
             # Reload the updated doc back into current_engine
             with open(docx_path, "rb") as f:
                 current_engine.load_document(io.BytesIO(f.read()))
                 
             return jsonify({
                "message": "TOC and Page Numbers updated successfully.",
                "preview": current_engine.get_preview_data(),
                "html_preview": current_engine.get_html_preview()
             })
        else:
             return jsonify({"error": "Word engine failed to verify TOC updates."}), 500

    except Exception as e:
        logging.error(f"TOC Update Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/format_official', methods=['POST'])
def canvas_format_official():
    try:
        data = request.get_json()
        model_config = data.get("model_config", {})
        scope = data.get("scope", "all")
        processor = data.get("processor", "local")

        if not current_engine.staging_doc:
            current_engine.create_staging_copy()

        context = current_engine.get_preview_data()
        code = llm_engine.generate_formatting_code(context, model_config, scope, processor)
        
        success, error_msg = current_engine.execute_code(code)
        
        if not success:
             return jsonify({"error": f"Failed to execute formatting code: {error_msg}"}), 400

        return jsonify({
            "message": "Formatting applied", 
            "code_executed": code,
            "preview": current_engine.get_preview_data(),
            "html_preview": current_engine.get_html_preview(),
            "is_staging": True
        })
    except Exception as e:
        logging.error(f"Canvas Format Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/modify_local', methods=['POST'])
def canvas_modify_local():
    try:
        data = request.get_json()
        file_path = data.get("file_path")
        instruction = data.get("instruction")
        
        if not file_path or not instruction:
            return jsonify({"error": "file_path and instruction are required"}), 400
        
        current_engine.load_from_path(file_path)
        context = current_engine.get_preview_data()
        code = llm_engine.generate_code(instruction, context)
        
        success, error_msg = current_engine.execute_code(code)
        if not success:
             return jsonify({"error": f"Failed to execute AI code: {error_msg}"}), 500
        
        current_engine.save_to_path(file_path)
        
        return jsonify({
            "message": "File processed and saved",
            "file_path": file_path,
            "preview": current_engine.get_preview_data(),
             "html_preview": current_engine.get_html_preview()
        })
    except Exception as e:
        logging.error(f"Canvas Modify Local Error: {e}")
        return jsonify({"error": str(e)}), 500

@canvas_bp.route('/export_docx', methods=['POST'])
def export_docs():
    data = request.json
    markdown_content = data.get('markdown', '')

    try:
        doc = Document()
        
        # Setup Page Margins
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        def set_font(run, font_name, size_pt, bold=False):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.bold = bold

        lines = markdown_content.split('\n')
        first_line_processed = False

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
            
            # Check for image: ![alt](/static/images/uuid.png)
            img_match = re.search(r'!\[.*?\]\((.*?)\)', stripped_line)
            if img_match:
                img_url = img_match.group(1)
                if '/static/images/' in img_url:
                    filename = img_url.split('/static/images/')[-1]
                    img_path = os.path.join(current_app.root_path, 'static', 'images', filename)
                    
                    if os.path.exists(img_path):
                        try:
                            doc.add_picture(img_path, width=Cm(15))
                            last_p = doc.paragraphs[-1] 
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as img_err:
                            logging.error(f"Failed to add image {img_path}: {img_err}")
                            p = doc.add_paragraph(f"[图片加载失败: {filename}]")
                    else:
                        p = doc.add_paragraph(f"[图片丢失: {filename}]")
                continue

            clean_text = stripped_line.replace('*', '').replace('#', '').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28) 

            if not first_line_processed:
                # TITLE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(clean_text)
                set_font(run, '小标宋体', 22, bold=False)
                first_line_processed = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                is_l1 = re.match(r'^[一二三四五六七八九十]+、', clean_text)
                is_l2 = re.match(r'^（[一二三四五六七八九十]+）', clean_text)
                
                if is_l1:
                    run = p.add_run(clean_text)
                    set_font(run, '黑体', 16)
                elif is_l2:
                    run = p.add_run(clean_text)
                    set_font(run, '楷体_GB2312', 16)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.1) 
                    run = p.add_run(clean_text)
                    set_font(run, '仿宋_GB2312', 16)

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return send_file(
            f, 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="smart_export.docx"
        )
    except Exception as e:
        logging.error(f"Smart Canvas Export Error: {e}")
        return jsonify({"error": str(e)}), 500
