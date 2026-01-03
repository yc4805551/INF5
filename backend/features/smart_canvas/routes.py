from flask import Blueprint, request, jsonify, current_app
import logging
import os
from core.services import current_engine

smart_canvas_bp = Blueprint('smart_canvas', __name__)

@smart_canvas_bp.route('/upload', methods=['POST'])
def smart_canvas_upload():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        # Save uploaded file temporarily
        temp_path = os.path.join(current_app.root_path, 'temp_upload.docx')
        file.save(temp_path)
        
        # Image output directory
        images_dir = os.path.join(current_app.root_path, 'static', 'images')
        
        # Extract using mammoth (via current_engine helper)
        markdown_content = current_engine.extract_with_images(temp_path, images_dir)
        
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
        return jsonify({
            "message": "File processed",
            "markdown": markdown_content
        })
    except Exception as e:
        logging.error(f"Smart Canvas Upload Error: {e}")
        return jsonify({"error": str(e)}), 500

@smart_canvas_bp.route('/transfer_to_canvas', methods=['POST'])
def transfer_to_canvas():
    try:
        data = request.json
        markdown_content = data.get('markdown', '')
        
        # 1. Convert Markdown to Docx (Simplified logic similar to canvas export)
        # We need to import necessary modules locally to avoid circular imports or pollution if they aren't at top
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import io
        import re
        
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        def set_font(run, font_name, size_pt, bold=False):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.bold = bold

        lines = markdown_content.split('\n')
        first_line_processed = False

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
            
            # Simple Image Skip (or handle if we want)
            if stripped_line.startswith('!['):
                continue

            clean_text = stripped_line.replace('*', '').replace('#', '').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28)

            if not first_line_processed:
                # Assume first line is Title
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(clean_text)
                set_font(run, '方正小标宋简体', 22, bold=False)
                first_line_processed = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Pattern matching for Headings
                is_l1 = re.match(r'^[一二三四五六七八九十]+、', clean_text)
                is_l2 = re.match(r'^（[一二三四五六七八九十]+）', clean_text)
                
                if is_l1:
                    run = p.add_run(clean_text)
                    set_font(run, '黑体', 16)
                elif is_l2:
                    run = p.add_run(clean_text)
                    set_font(run, '楷体_GB2312', 16)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.1)
                    run = p.add_run(clean_text)
                    set_font(run, '仿宋_GB2312', 16)

        # 2. Save to stream
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        # 3. Load into Main Engine
        current_engine.load_document(f)
        
        # 4. Return preview to confirm
        return jsonify({
            "message": "Transfer successful",
            "preview": current_engine.get_preview_data(limit=100),
            "html_preview": current_engine.get_html_preview(limit=100)
        })

    except Exception as e:
        logging.error(f"Transfer Error: {e}")
        return jsonify({"error": str(e)}), 500

@smart_canvas_bp.route('/export', methods=['POST'])
def smart_canvas_export():
    try:
        data = request.json
        markdown_content = data.get('markdown', '')
        
        # 1. Convert Markdown to Docx (Same logic as transfer_to_canvas)
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import io
        import re
        from flask import send_file
        
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        def set_font(run, font_name, size_pt, bold=False):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.bold = bold

        lines = markdown_content.split('\n')
        first_line_processed = False

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue
            
            # Simple Image Skip (or handle if we want)
            if stripped_line.startswith('!['):
                continue

            clean_text = stripped_line.replace('*', '').replace('#', '').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28)

            if not first_line_processed:
                # Assume first line is Title
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(clean_text)
                set_font(run, '方正小标宋简体', 22, bold=False)
                first_line_processed = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Pattern matching for Headings
                is_l1 = re.match(r'^[一二三四五六七八九十]+、', clean_text)
                is_l2 = re.match(r'^（[一二三四五六七八九十]+）', clean_text)
                
                if is_l1:
                    run = p.add_run(clean_text)
                    set_font(run, '黑体', 16)
                elif is_l2:
                    run = p.add_run(clean_text)
                    set_font(run, '楷体_GB2312', 16)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.1)
                    run = p.add_run(clean_text)
                    set_font(run, '仿宋_GB2312', 16)

        # 2. Save to stream
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return send_file(f, as_attachment=True, download_name='co_creation_export.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        logging.error(f"Export Error: {e}")
        return jsonify({"error": str(e)}), 500
