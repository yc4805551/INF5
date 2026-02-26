
import os
import io
import json
import logging
import traceback
import tempfile
import pandas as pd
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
from PIL import Image

# Core LLM Engine for OCR
from core.llm_engine import LLMEngine
from .config import OCR_MODEL_PROVIDER, OCR_MODEL_NAME, OCR_API_KEY, OCR_ENDPOINT

logger = logging.getLogger(__name__)

class SmartFileAgent:
    def __init__(self, use_llm_clean=False, cleaning_model_config=None, ocr_provider=None):
        self.use_llm_clean = use_llm_clean
        self.cleaning_model_config = cleaning_model_config
        self.merged_buffer = []

        # Resolve OCR Configuration
        self.ocr_provider = ocr_provider or OCR_MODEL_PROVIDER
        
        from core.llm_config import get_llm_config_manager
        cm = get_llm_config_manager()
        provider_config = cm.get_provider_config(self.ocr_provider)

        # Priority 1: Direct OCR_* Env Overrides (Dedicated Mode) 
        # Only use these overrides if they are explicitly configured AND we are using the default/env provider
        if OCR_ENDPOINT and self.ocr_provider == OCR_MODEL_PROVIDER and getattr(self, 'use_legacy_ocr', False):
             self.ocr_api_key = OCR_API_KEY
             self.ocr_model_name = OCR_MODEL_NAME
             self.ocr_endpoint = OCR_ENDPOINT
             self.ocr_provider = "custom"
        else:
             # Priority 2: Provider Lookup (via LLMConfigManager)
             self.ocr_api_key = provider_config.get("apiKey")
             self.ocr_model_name = provider_config.get("model")
             self.ocr_endpoint = provider_config.get("endpoint")

        # Initialize OCR Engine independently
        self.ocr_engine = LLMEngine(api_key=self.ocr_api_key)

    def process_files(self, files_data):
        """
        Generator function to yield progress logs and final result.
        files_data: List of dicts {'filename': str, 'content': bytes}
        """
        total_files = len(files_data)
        # Sort files by name to ensure consistent order
        files_data.sort(key=lambda x: x['filename'])
        
        yield json.dumps({"type": "log", "message": f"Found {total_files} files to process..."}) + "\n"
        yield json.dumps({"type": "log", "message": f"Active OCR Configuration: Provider=[{self.ocr_provider}], Model=[{self.ocr_model_name}]"}) + "\n"

        for index, file_info in enumerate(files_data):
            file_name = file_info['filename']
            file_content = file_info['content']
            
            yield json.dumps({"type": "log", "message": f"Processing [{index+1}/{total_files}]: {file_name}..."}) + "\n"
            
            try:
                # Content is already bytes
                file_bytes = io.BytesIO(file_content)
                
                # 2. Route Processing (Pass in-memory BytesIO)
                processed_text = ""
                ext = os.path.splitext(file_name)[1].lower()

                if ext in ['.xlsx', '.xls']:
                    processed_text = self._process_excel(file_bytes)
                elif ext == '.docx':
                    processed_text = self._process_word(file_bytes)
                elif ext == '.pdf':
                     processed_text = ""
                     for update in self._process_pdf_smart(file_content, file_name):
                         if isinstance(update, dict):
                             processed_text = update.get("text", "")
                         else:
                             yield update
                elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
                     yield json.dumps({"type": "log", "message": f"  - [{file_name}] Identifying image with {self.ocr_model_name}..."}) + "\n"
                     processed_text = self._process_image(file_content)
                elif ext in ['.txt', '.md', '.csv']:
                    try:
                        processed_text = file_content.decode('utf-8')
                    except UnicodeDecodeError:
                        processed_text = file_content.decode('gbk', errors='ignore')

                else:
                    processed_text = f"[Skipped unsupported file: {file_name}]"

                # 3. Format and Append
                formatted_section = f"\n\n{'='*20}\n文件名: {file_name}\n{'='*20}\n\n{processed_text}"
                self.merged_buffer.append(formatted_section)
                
            except Exception as e:
                error_msg = f"Error processing {file_name}: {str(e)}"
                logger.error(error_msg)
                logger.error(traceback.format_exc())
                yield json.dumps({"type": "log", "message": f"  - [Error] {error_msg}"}) + "\n"
                self.merged_buffer.append(f"\n\n[Error processing {file_name}]\n")

        # 4. Merger
        yield json.dumps({"type": "log", "message": "Merging results..."}) + "\n"
        full_text = "".join(self.merged_buffer)

        # 5. Optional: LLM Cleaning
        if self.use_llm_clean and self.cleaning_model_config:
             pass

        # Final Yield
        yield json.dumps({"type": "result", "text": full_text}) + "\n"
        yield json.dumps({"type": "log", "message": "Done."}) + "\n"

    def _process_excel(self, file_stream):
        try:
            # pandas read_excel supports file-like objects (BytesIO)
            df = pd.read_excel(file_stream)
            return df.to_markdown(index=False)
        except Exception as e:
            return f"[Excel Error: {str(e)}]"

    def _process_word(self, file_stream):
        try:
            # python-docx supports file-like objects (BytesIO)
            doc = Document(file_stream)
            full_text = []
            
            # 1. Extract Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 2. Extract Tables (Append them to ensure data isn't lost)
            if doc.tables:
                full_text.append("\n\n--- Tables Extracted from Document ---\n")
                for table in doc.tables:
                    # Simple Markdown Table Converter
                    # Get all rows
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        rows.append(cells)
                    
                    if not rows:
                        continue

                    # Determine columns (max length of any row)
                    # Some merged cells might cause issues, this is a basic approximation
                    num_cols = max(len(r) for r in rows) if rows else 0
                    
                    if num_cols > 0:
                        # 1. Header
                        header = rows[0]
                        # Pad header if needed
                        header += [''] * (num_cols - len(header))
                        full_text.append("| " + " | ".join(header) + " |")
                        
                        # 2. Separator
                        full_text.append("| " + " | ".join(['---'] * num_cols) + " |")
                        
                        # 3. Body
                        for row in rows[1:]:
                            # Pad row if needed
                            row += [''] * (num_cols - len(row))
                            full_text.append("| " + " | ".join(row) + " |")
                        
                        full_text.append("\n") # Spacing between tables

            return "\n".join(full_text)
        except Exception as e:
            return f"[Word Error: {str(e)}]"

    def _process_pdf_smart(self, file_bytes, file_name):
        try:
            import fitz
            import re
            import concurrent.futures
            
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            results = [None] * len(doc)
            ocr_prompt = (
                "你是一个专业的中文文档和表格排版专家。请将图片中的内容精准地转换为 Markdown 格式。\n"
                "要求：\n"
                "1. 请修正可能由于扫描造成的错别字或折叠。\n"
                "2. 如果图片中包含表格，请务必使用标准的 Markdown 表格语法 ('|---|') 进行严谨的还原，不要漏掉合并单元格或表头。\n"
                "3. 不要输出任何开场白或解释文字，直接输出转换后的 Markdown。\n"
                "4. 如果图片内容完全无法辨认、或者包含大量无意义的乱码和符号，请直接输出 '[UNREADABLE]'，不要强行编造或输出乱码。"
            )
            
            # Using ThreadPoolExecutor to run OCR concurrently.
            # Max_workers=5 is a safe threshold for SiliconFlow's concurrent rate limits.
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                futures = {}
                
                for i, page in enumerate(doc):
                    page_text = page.get_text()
                    
                    # Heuristic to detect gibberish per page
                    non_ws_text = re.sub(r'\s+', '', page_text)
                    valid_chars = re.findall(r'[\u4e00-\u9fa5A-Za-z0-9]', non_ws_text)
                    chinese_chars = re.findall(r'[\u4e00-\u9fa5]', non_ws_text)
                    
                    is_gibberish = False
                    if len(non_ws_text) > 0:
                        if (len(valid_chars) / len(non_ws_text)) < 0.2:
                            is_gibberish = True
                        elif len(chinese_chars) < 5 and len(valid_chars) > 50:
                            is_gibberish = True
                    
                    if len(page_text) < 20 or is_gibberish:
                        yield json.dumps({"type": "log", "message": f"  - [{file_name}] 第 {i+1} 页启用 Tesseract虚拟切片 混合识别..."}) + "\n"
                        # Extract the image up front so we can pass bytes to the thread (avoiding PyMuPDF thread-safety issues with doc)
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        
                        future = executor.submit(self._process_page_virtual_slicing, img_data, ocr_prompt)
                        futures[future] = i
                    else:
                        results[i] = page_text
                
                # As each page finishes OCR, yield progress to the UI
                for future in concurrent.futures.as_completed(futures):
                    page_idx = futures[future]
                    try:
                        vision_response = future.result()
                        results[page_idx] = vision_response
                        yield json.dumps({"type": "log", "message": f"  - [{file_name}] 第 {page_idx+1} 页 OCR 识别完成！"}) + "\n"
                    except Exception as e:
                        results[page_idx] = f"[OCR Error on page {page_idx+1}: {e}]"
                        yield json.dumps({"type": "log", "message": f"  - [{file_name}] 第 {page_idx+1} 页 OCR 识别失败：{e}"}) + "\n"
            
            doc.close()
            # Join the results precisely in page order
            yield {"text": "\n\n".join(filter(None, results))}
        except Exception as e:
            logger.error(f"PDF Error: {e}")
            yield {"text": f"[PDF Extract Error: {e}]"}

    def _scrub_ghosts(self, text):
        import re
        # "[UNREADABLE]" is deliberately removed from this list because it's our requested output for 
        # completely dead spaces. It shouldn't trigger a verbose error message replacing valid text.
        ghost_signatures = ["畜牧兽医", "<|LOC_", "omoData", "阴夜雨", "重夜雨", "لن قم المو", "Employee"]
        if any(ghost in text for ghost in ghost_signatures):
            return "" # Silently drop the hallucinated text
        
        # Scrub generative VLM token loops (e.g. }}}}}}} or 1.1.1.1.1.)
        text = re.sub(r'(\}\s*){5,}', '', text)
        text = re.sub(r'(\{\s*){5,}', '', text)
        text = re.sub(r'(1\.\s*){5,}', '', text)
        text = re.sub(r'(-\s*){10,}', '', text)
        text = re.sub(r'(\]\s*){5,}', '', text)
        text = re.sub(r'(\\\s*\]\s*){5,}', '', text)
        
        # Aggressive filter for DeepSeek-OCR / PaddleOCR table artifacts
        # e.g. \end{array}\right\} repeated 10 times
        text = re.sub(r'(\\end\{array\}\\right\\\}\$\s*){3,}', '', text) 
        text = re.sub(r'(\\begin\{array\}\\right\\\}\$\s*){3,}', '', text)
        
        # The Ultimate Wildcard Filter:
        # Matches ANY sequence of 1 to 50 characters that repeats identically at least 8 times in a row.
        # This completely wipes out generative model infinite token loops (like `}}]}}]}}]}}`) 
        # without hardcoding the exact symbols.
        text = re.sub(r'(.{1,50}?)\1{8,}', '', text)
        
        # Strip out localized [UNREADABLE] markers so they don't clutter the final markdown
        return text.replace("[UNREADABLE]", "").strip()
    def _is_image_mostly_blank(self, pil_img, min_pixel_threshold=240, max_color_diff=15):
        try:
            from PIL import Image, ImageStat
            # Convert to grayscale to check brightness and variance
            gray = pil_img.convert('L')
            
            # Low-pass filter: resize to a tiny thumbnail using Bilinear interpolation.
            # This completely annihilates microscopic scanner dust/noise (averaging it out to white)
            # but preserves the larger structural gray bands created by actual text.
            tiny = gray.resize((100, 100), Image.Resampling.BILINEAR)
            
            stat = ImageStat.Stat(tiny)
            min_val, max_val = stat.extrema[0]
            
            # If the darkest pixel is very bright white/gray (e.g. >= 240)
            if min_val >= min_pixel_threshold:
                return True
                
            # Or if there is almost no contrast (difference between darkest and lightest pixel is tiny)
            # This catches solid gray/black/colored boxes that have no text or features
            if (max_val - min_val) < max_color_diff:
                return True
                
            return False
        except Exception:
            return False

    def _filter_headers_footers(self, blocks, img_height, margin_percent=0.08):
        """Scalpel 1: Remove blocks that appear in the extreme top/bottom margins."""
        clean = []
        top_threshold = img_height * margin_percent
        bottom_threshold = img_height * (1.0 - margin_percent)
        for b in blocks:
            y_center = b['y'] + (b['h'] / 2)
            if top_threshold < y_center < bottom_threshold:
                clean.append(b)
        return clean

    def _cluster_x_axis(self, blocks):
        """Scalpel 2: Detect and sort multi-column layouts based on X-coordinates."""
        if not blocks:
            return []
        
        # Simple clustering: if the centers of blocks fall cleanly into 
        # left-half vs right-half, they are likely columns.
        # A more robust solution for academic papers uses K-Means or Histograms. 
        # Here we do a 2-column heuristic check based on page width.
        
        # Get overall page bounds from remaining blocks
        min_x = min(b['x'] for b in blocks)
        max_x = max(b['x'] + b['w'] for b in blocks)
        page_width = max_x - min_x
        mid_point = min_x + (page_width / 2)
        
        col_left = []
        col_right = []
        
        for b in blocks:
            x_center = b['x'] + (b['w'] / 2)
            if x_center < mid_point:
                col_left.append(b)
            else:
                col_right.append(b)
                
        # Sort each column individually top-to-bottom
        col_left.sort(key=lambda x: x['y'])
        col_right.sort(key=lambda x: x['y'])
        
        return col_left + col_right # Read left column entirely, then right column

    def _stitch_y_axis(self, blocks, max_line_gap_ratio=1.5):
        """Scalpel 3: Reassemble paragraphs by calculating Y-gaps between consecutive sorted blocks."""
        if not blocks:
            return []
            
        paragraphs = []
        current_para = [blocks[0]]
        
        for i in range(1, len(blocks)):
            prev = blocks[i-1]
            curr = blocks[i]
            
            # Distance from bottom of prev to top of curr
            delta_y = curr['y'] - (prev['y'] + prev['h'])
            avg_height = (prev['h'] + curr['h']) / 2
            
            # If the gap is small (like a normal line break), it belongs to the same paragraph
            # Allow delta_y to be slightly negative for overlapping lines or slight tilts
            if delta_y < (avg_height * max_line_gap_ratio):
                current_para.append(curr)
            else:
                # Gap is too big, start a new paragraph
                paragraphs.append(current_para)
                current_para = [curr]
                
        if current_para:
            paragraphs.append(current_para)
            
        return paragraphs

    def _auto_crop_whitespace(self, pil_img):
        try:
            import numpy as np
            from PIL import Image
            
            # Convert to grayscale to evaluate pixel intensity
            gray = pil_img.convert('L')
            arr = np.array(gray)
            
            # Threshold: pixels darker than 200 are considered 'ink'
            ink_threshold = 200
            
            # Count how many ink pixels are in each row
            ink_pixels_per_row = np.sum(arr < ink_threshold, axis=1)
            
            # A row is considered to have 'text' if it has at least 10 ink pixels 
            # (This ignores single specks of dust, staple marks, or scanner noise)
            text_rows = np.where(ink_pixels_per_row > 10)[0]
            
            if len(text_rows) > 0:
                top = text_rows[0]
                bottom = text_rows[-1]
                
                # Do the same for columns, but only within the y-bounds we just found
                ink_pixels_per_col = np.sum(arr[top:bottom, :] < ink_threshold, axis=0)
                text_cols = np.where(ink_pixels_per_col > 10)[0]
                
                if len(text_cols) > 0:
                    left = text_cols[0]
                    right = text_cols[-1]
                    
                    # Add a comfortable margin of 40 pixels so descenders aren't clipped
                    margin = 40
                    h, w = arr.shape
                    top = max(0, int(top - margin))
                    bottom = min(h, int(bottom + margin))
                    left = max(0, int(left - margin))
                    right = min(w, int(right + margin))
                    
                    # If the cropped area is extremely tiny (like a single smudge), don't crop
                    if (bottom - top) < 50 or (right - left) < 50:
                        return pil_img
                        
                    return pil_img.crop((left, top, right, bottom))
            return pil_img
        except Exception as e:
            return pil_img

    def _process_page_virtual_slicing(self, img_bytes, prmpt):
        try:
            import pytesseract
            from pytesseract import Output
            import cv2
            import numpy as np
            import os
            
            # Setup environment Tesseract Path if provided (For remote deployments)
            tess_path = os.getenv("VITE_TESSERACT_PATH")
            if tess_path and os.path.exists(tess_path):
                pytesseract.pytesseract.tesseract_cmd = tess_path
            
            # 1. Convert Img Bytes to OpenCV format
            nparr = np.frombuffer(img_bytes, np.uint8)
            cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            img_height, img_width = cv_img.shape[:2]
            
            # 2. Call Tesseract for "Full Radar Scan" (Fast local OCR)
            scan_data = pytesseract.image_to_data(cv_img, lang='chi_sim', output_type=Output.DICT)
            
            blocks = []
            for i in range(len(scan_data['text'])):
                conf = scan_data['conf'][i]
                text = scan_data['text'][i].strip()
                # Accept conf > 10 to catch small texts, or any valid string if conf fails parsing
                if (isinstance(conf, str) and conf != '-1') or (isinstance(conf, (int, float)) and conf > 10):
                    if text:
                        blocks.append({
                            'text': text,
                            'x': scan_data['left'][i],
                            'y': scan_data['top'][i],
                            'w': scan_data['width'][i],
                            'h': scan_data['height'][i]
                        })
            
            # If Tesseract completely fails to find anything, fallback to whole-page vision API
            if not blocks:
                return self._slice_and_ocr_image(img_bytes, prmpt)
                
            # 3. Apply the Three Scalpels (Pure Math Spatial Clustering)
            blocks = self._filter_headers_footers(blocks, img_height) # Scalpel 1: Y-axis Header cut
            columns = self._cluster_x_axis(blocks)                    # Scalpel 2: X-axis Columns
            paragraphs = self._stitch_y_axis(columns)                 # Scalpel 3: Y-axis Stitching
            
            final_markdown = ""
            for para in paragraphs:
                if not para:
                    continue
                    
                # Calculate the bounding box for this entire stitched paragraph
                min_x = min(b['x'] for b in para)
                min_y = min(b['y'] for b in para)
                max_x = max(b['x'] + b['w'] for b in para)
                max_y = max(b['y'] + b['h'] for b in para)
                
                # Heuristic for Table Anomaly Detection:
                # If a "paragraph" contains a massive amount of strictly separated small blocks
                # grouped tightly together in a wide array, it's likely a complex graphical table.
                # Thresholds: More than 15 independent blocks in a single stitched paragraph
                # AND total width is quite wide (> 40% of page width).
                para_width = max_x - min_x
                if len(para) > 15 and para_width > (img_width * 0.4):
                    # HYBRID ROUTING: Snip out the anomaly and send ONLY this tight bounding box to DeepSeek API
                    # Add 10px margin padding
                    p_min_y = max(0, min_y - 10)
                    p_max_y = min(img_height, max_y + 10)
                    p_min_x = max(0, min_x - 10)
                    p_max_x = min(img_width, max_x + 10)
                    
                    table_snipe = cv_img[p_min_y:p_max_y, p_min_x:p_max_x]
                    
                    # Convert cv2 image back to bytes for the API call
                    success, buffer = cv2.imencode('.png', table_snipe)
                    if success:
                        table_bytes = buffer.tobytes()
                        # Call DeepSeek/Paddle API
                        vlm_result = self._call_vision_api(table_bytes, "这是一个表格局部截图，严格输出Markdown表格。不要解释。")
                        final_markdown += vlm_result + "\n\n"
                    else:
                        # Fallback to Tesseract chunks if encoding fails
                        final_markdown += " ".join([b['text'] for b in para]) + "\n\n"
                else:
                    # Normal Text Paragraph: Use the instant, free Tesseract extraction
                    para_text = ""
                    for i, b in enumerate(para):
                        if i > 0:
                            # If there's a tiny x-gap between words within the same para, add a little space
                            prev = para[i-1]
                            if b['x'] - (prev['x'] + prev['w']) > b['h'] * 0.5:
                                para_text += " "
                        para_text += b['text']
                    
                    final_markdown += para_text + "\n\n"
            
            return final_markdown
            
        except ImportError as ie:
            # If Remote server hasn't installed pytesseract/opencv, gracefully fallback to physical slice mode
            import logging
            logging.warning(f"Virtual Slicing Dependencies missing, falling back to basic OCR: {ie}")
            return self._slice_and_ocr_image(img_bytes, prmpt)
        except Exception as e:
            import logging
            logging.error(f"Virtual Slicing failed: {e}")
            # Fallback on any mathematical error
            return self._slice_and_ocr_image(img_bytes, prmpt)

    def _slice_and_ocr_image(self, img_bytes, prompt, max_pixels=4000000, max_width=1600):
        try:
            from PIL import Image
            import io
            import base64
            
            img = Image.open(io.BytesIO(img_bytes))
            if img.mode != 'RGB':
                img = img.convert('RGB')
                
            # Smart Auto-Crop: Remove extreme whitespace from document borders (especially the bottom)
            # This absolutely prevents "Prompt Bleeding" / "Instruction Hallucination" from VLMs like PaddleOCR
            img = self._auto_crop_whitespace(img)
                
            orig_width, orig_height = img.size
            # Cap the maximum width to 1600 pixels (very high res) to avoid excessive slice counts on 4K/8K scans
            if orig_width > max_width:
                new_width = max_width
                new_height = int((max_width / orig_width) * orig_height)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
            width, height = img.size
            if width * height <= max_pixels:
                if self._is_image_mostly_blank(img):
                    return "" # Skip entirely blank images before OCR to prevent hallucinations
                
                # Image is small enough, process as usual but convert to PNG to preserve sharp text edges
                # (Generative VLMs often hallucinate on JPEG compression artifacts around text)
                buffered = io.BytesIO()
                img.save(buffered, format="PNG")
                b64_img = "data:image/png;base64," + base64.b64encode(buffered.getvalue()).decode('utf-8')
                return self._scrub_ghosts(self._call_vision_api(b64_img, prompt))
                
            # Image is too large. Slice it horizontally to preserve resolution.
            num_slices = (width * height // max_pixels) + 1
            slice_height = height // num_slices
            
            full_text = [None] * num_slices
            
            def process_slice(i):
                top = i * slice_height
                bottom = height if i == num_slices - 1 else (i + 1) * slice_height
                
                # Add a 60-pixel overlap to avoid cutting text lines in half completely
                overlap = 60 if i > 0 else 0
                box = (0, max(0, top - overlap), width, bottom)
                
                slice_img = img.crop(box)
                
                if self._is_image_mostly_blank(slice_img):
                    return "" # Skip this slice if it's completely blank
                
                buffered = io.BytesIO()
                slice_img.save(buffered, format="PNG")
                b64_img = "data:image/png;base64," + base64.b64encode(buffered.getvalue()).decode('utf-8')
                
                slice_prompt = prompt
                if num_slices > 1:
                    slice_prompt += f"\n\n（注意：这是超长图的第 {i+1}/{num_slices} 部分截图，请直接输出图里的内容文本，由于可能截断了部分图形不要擅自发散废话。如果由于切片导致本片完全是空白或者乱码，请仅输出 [UNREADABLE]）"
                
                response = self._call_vision_api(b64_img, slice_prompt)
                return self._scrub_ghosts(response)

            import concurrent.futures
            # Process slices concurrently (up to 3 at a time to avoid heavy rate limits if called within a loop)
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                futures = {executor.submit(process_slice, i): i for i in range(num_slices)}
                for future in concurrent.futures.as_completed(futures):
                    i = futures[future]
                    try:
                        res = future.result()
                        full_text[i] = res
                    except Exception as e:
                        import logging
                        logging.error(f"Image slice {i} error: {e}")
                        full_text[i] = ""
                
            return "\n\n".join(filter(None, full_text))
        except Exception as e:
            import base64
            import logging
            logging.error(f"Image slice processing error: {e}")
            # Fallback to direct call, hoping for the best
            b64_img = "data:image/png;base64," + base64.b64encode(img_bytes).decode('utf-8')
            return self._scrub_ghosts(self._call_vision_api(b64_img, prompt))



    def _process_image(self, file_bytes):
        try:
            prompt = (
                "你是一个专业的中文文档和表格排版专家。请将图片中的内容精准地转换为 Markdown 格式。\n"
                "要求：\n"
                "1. 请修正可能由于扫描造成的错别字或折叠。\n"
                "2. 如果图片中包含表格，请务必使用标准的 Markdown 表格语法 ('|---|') 进行严谨的还原，不要漏掉合并单元格或表头。\n"
                "3. 不要输出任何开场白或解释文字，直接输出转换后的 Markdown。\n"
                "4. 如果图片内容完全无法辨认、或者包含大量无意义的乱码和符号，请直接输出 '[UNREADABLE]'，不要强行编造或输出乱码。"
            )
            return self._slice_and_ocr_image(file_bytes, prompt)
        except Exception as e:
            return f"[Image Processing Error: {str(e)}]"

    def _call_vision_api(self, b64_image, prompt):
        # Implementation of OpenAI-compatible Vision request
        # Assuming 'DeepSeek OCR' or 'depOCR' creates an OpenAI-compatible endpoint that accepts image_url
        
        import requests
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.ocr_api_key}"
        }
        
        content_payload = [
            {"type": "text", "text": prompt},
            {
                "type": "image_url",
                "image_url": {
                    "url": b64_image
                }
            }
        ]

        payload = {
            "model": self.ocr_model_name,
            "messages": [
                {
                    "role": "user",
                    "content": content_payload
                }
            ],
            "max_tokens": 4096,
            "temperature": 0.2, # Slightly > 0 to prevent infinite loops on degenerate inputs
            "top_p": 0.9,
            "presence_penalty": 0.5, # Prevent endless repetition of characters like `}}}}` or `1.1.`
            "frequency_penalty": 0.5
        }
        
        # Use pre-resolved endpoint
        endpoint = self.ocr_endpoint
        
        if not endpoint:
            # Fallback if somehow missing (shouldn't happen given __init__ logic)
            return "[Error: No Endpoint Configured for OCR Model]"
        
        if not endpoint:
            return "[Error: No Endpoint Configured for OCR Model]"

        if "/chat/completions" not in endpoint:
             endpoint = endpoint.rstrip("/") + "/chat/completions"

        import time
        max_retries = 2
        for attempt in range(max_retries):
            try:
                # Shorter 30-second timeout. Generative VLMs can sometimes hang indefinitely on bad inputs.
                # If it takes >30s, it's better to fail fast and retry/skip rather than deadlocking the ThreadPool.
                response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
                if response.status_code == 200:
                    res_json = response.json()
                    return res_json['choices'][0]['message']['content']
                elif response.status_code == 429: # Too Many Requests
                    if attempt < max_retries - 1:
                        time.sleep(2 * (attempt + 1)) # Exponential backoff
                        continue
                    return f"[OCR API Rate Limit (429): Please try again later.]"
                else:
                    return f"[OCR API Error {response.status_code}: {response.text}]"
            except requests.exceptions.Timeout:
                if attempt < max_retries - 1:
                    logger.warning(f"OCR API Timeout, retrying... ({attempt+1}/{max_retries})")
                    continue
                return f"[OCR Request Timed Out after 30s]"
            except Exception as e:
                return f"[OCR Request Failed: {str(e)}]"
