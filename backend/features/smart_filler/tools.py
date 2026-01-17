import logging
import pandas as pd
from core.services import current_engine

import logging
import os
import datetime

logger = logging.getLogger(__name__)

def log_debug(message):
    """Writes to smart_filler_debug.log in backend root."""
    try:
        # tools.py is in backend/features/smart_filler/
        # 1. smart_filler, 2. features, 3. backend
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        log_path = os.path.join(base_dir, "smart_filler_debug.log")
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"[{timestamp}] {message}\n")
    except Exception:
        pass # Fallback


class ToolRegistry:
    def __init__(self, service_context):
        """
        service_context: The SmartFillerService instance (to access current_df, etc.)
        """
        self.context = service_context
        self.tools_map = {
            "read_excel_summary": self.read_excel_summary,
            "find_anchor_in_word": self.find_anchor_in_word,
            "write_word_content": self.write_word_content,
            "read_document_structure": self.read_document_structure,
            "read_source_content": self.read_source_content,
            "copy_image_to_word": self.copy_image_to_word,
            "execute_document_script": self.execute_document_script
        }
        
        # Schema Definitions (MCP-inspired)
        self.definitions = [
            {
                "name": "read_excel_summary",
                "description": "Returns a summary of the loaded Excel file, including columns, row count, and a preview of the first 3 rows.",
                "input_schema": {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            },
            {
                "name": "read_source_content",
                "description": "Reads the text content of reduced source documents (Reference Docx/Images). Use this to understand the source material.",
                "input_schema": {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            },
            {
                "name": "find_anchor_in_word",
                "description": "Searches for a specific text string in the Word document to locate its position (Table/Row/Cell or Paragraph).",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "text": {"type": "string", "description": "The exact text to search for."}
                    },
                    "required": ["text"]
                }
            },
            {
                "name": "write_word_content",
                "description": "Writes text into the Word document at a specific location.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "location": {"type": "string", "description": "Location string returned by find_anchor or similar (e.g., 'Table 0, Row 1, Cell 2' or 'Paragraph 5')."},
                        "text": {"type": "string", "description": "The text content to write."},
                        "relative": {"type": "string", "enum": ["right", "replace", "append"], "description": "How to write relative to location. 'right' (default) puts in next cell. 'replace' overwrites. 'append' adds to end."}
                    },
                    "required": ["location", "text"]
                }
            },
            {
                "name": "read_document_structure",
                "description": "Extracts the outline (headers) or structure of the target Word document.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "mode": {"type": "string", "enum": ["outline"], "default": "outline"}
                    },
                    "required": []
                }
            },
            {
                "name": "copy_image_to_word",
                "description": "Inserts the currently loaded source image into the Word document.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "location": {"type": "string", "description": "Location to insert image (e.g. 'Table 0, Row 0, Cell 1' or 'Paragraph 2')."},
                        "width_inches": {"type": "number", "default": 2.0, "description": "Width of image in inches."}
                    },
                    "required": ["location"]
                }
            },
            {
                "name": "execute_document_script",
                "description": "Advanced: Executes a Python script to manipulate the document. Use 'doc' (python-docx Document object) and 'pd' (pandas) in your script.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "script_code": {"type": "string", "description": "Valid Python code to execute."}
                    },
                    "required": ["script_code"]
                }
            }
        ]

    def get_tool(self, name):
        return self.tools_map.get(name)
        
    def get_tools_schema(self):
        """Returns the list of tool definitions for the System Prompt."""
        return self.definitions

    def _ensure_save_path(self):
        """
        确保 current_engine 有有效的保存路径
        如果没有，自动创建缓存路径并保存文档
        
        这是防御性措施，避免 load_from_text 清空 original_path 导致写入失败
        """
        from core.services import current_engine
        import os
        
        if not current_engine.original_path:
            # 创建 .cache 目录
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            cache_dir = os.path.join(base_dir, ".cache")
            if not os.path.exists(cache_dir):
                os.makedirs(cache_dir, exist_ok=True)
            
            # 设置临时路径
            current_engine.original_path = os.path.join(cache_dir, "smart_filler_auto.docx")
            
            # 立即保存以建立文件
            if current_engine.doc:
                current_engine.doc.save(current_engine.original_path)
                self.context._log_debug(f"⚠️ Auto-created save path: {current_engine.original_path}")
            else:
                self.context._log_debug("⚠️ Warning: No document loaded, cannot auto-save")

    def read_excel_summary(self, **kwargs):
        """Tool: Returns summary of loaded Excel."""
        # Check internal SmartFiller context OR global references
        target_df = self.context.current_df
        
        if target_df is None:
            # Fallback: Check global references for any Excel File
            for ref in reversed(current_engine.reference_docs): # Use latest first
                if ref.get('type') == 'excel' and ref.get('df') is not None:
                    target_df = ref.get('df')
                    break
        
        if target_df is None:
            msg = "Error: No Excel file loaded."
            if self.context.current_context_text:
                msg += " However, a Docx/Text source IS loaded. You should probably use 'read_source_content' to read that instead."
            return msg
        columns = ", ".join(target_df.columns.tolist())
        rows = len(target_df)
        # Avoid huge dumps
        preview = target_df.head(3).to_dict(orient='records')
        return f"Excel loaded. Columns: [{columns}]. Total Rows: {rows}. Top 3 rows data: {preview}"

    def find_anchor_in_word(self, text, **kwargs):
        """Tool: Finds the location of text in Word."""
        target_doc = current_engine.staging_doc if current_engine.staging_doc else current_engine.doc
        if not target_doc:
            return "Error: No document loaded."
        
        # 1. Search Tables
        for t_idx, table in enumerate(target_doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if text in cell.text:
                        return f"Table {t_idx}, Row {r_idx}, Cell {c_idx}"
        
        # 2. Search Paragraphs
        for p_idx, para in enumerate(target_doc.paragraphs):
            if text in para.text:
                return f"Paragraph {p_idx}"
                
        return f"Text '{text}' not found."

    def write_word_content(self, location, text, relative="right", **kwargs):
        """Tool: Writes text to a location."""
        # 防御性检查：确保有保存路径
        self._ensure_save_path()
        
        target_doc = current_engine.staging_doc if current_engine.staging_doc else current_engine.doc
        if not target_doc:
            return "Error: No document loaded."

        try:
            if "Table" in location:
                # Parse "Table X, Row Y, Cell Z"
                parts = location.replace("Table ", "").replace("Row ", "").replace("Cell ", "").split(",")
                t_idx, r_idx, c_idx = int(parts[0]), int(parts[1]), int(parts[2])
                
                # Validate indices
                if t_idx >= len(target_doc.tables): return "Error: Table index out of range."
                table = target_doc.tables[t_idx]
                if r_idx >= len(table.rows): return "Error: Row index out of range."
                row = table.rows[r_idx]
                
                target_cell_idx = c_idx
                if relative == "right":
                    target_cell_idx += 1
                
                if target_cell_idx < len(row.cells):
                    row.cells[target_cell_idx].text = str(text)
                    if current_engine.original_path:
                        current_engine.doc.save(current_engine.original_path)
                    
                    # Signal Update
                    if hasattr(self.context, 'events'):
                        self.context.events.append({"type": "CANVAS_UPDATE"})
                        
                    return f"Success: Written '{text}' to Table {t_idx}, Row {r_idx}, Cell {target_cell_idx}"
                else:
                    return "Error: Target cell out of range."
                    
            elif "Paragraph" in location:
                p_idx = int(location.split(" ")[1])
                if p_idx >= len(target_doc.paragraphs): return "Error: Para index out of range."
                
                para = target_doc.paragraphs[p_idx]
                para.add_run(f" {text}")
                if current_engine.original_path:
                    target_doc.save(current_engine.original_path)
                    logger.info(f"Tool write saved to: {current_engine.original_path}")
                else:
                    logger.warning("Tool write failed to save: No original_path")
                
                # Signal Update
                if hasattr(self.context, 'events'):
                    self.context.events.append({"type": "CANVAS_UPDATE"})

                return f"Success: Appended '{text}' to Paragraph {p_idx}."
                
        except Exception as e:
            return f"Error writing content: {e}"

        return "Error: Invalid location format."

    def read_document_structure(self, file_type="docx", mode="outline", **kwargs):
        """
        Tool: Extracts document structure (Semantic Blocks).
        Currently a placeholder for the advanced DocxEngine capabilities.
        """
        if not current_engine.doc:
            return "Error: No document loaded."
            
        target_doc = current_engine.staging_doc if current_engine.staging_doc else current_engine.doc
            
        if mode == "outline":
            # Basic outline extraction using styles
            structure = []
            for i, para in enumerate(target_doc.paragraphs):
                if para.style.name.startswith('Heading'):
                    structure.append(f"Header: {para.text} (Para {i})")
            
            if not structure:
                return "No headers found in document."
            return "\n".join(structure)
            
        return "Unknown mode."

    def read_source_content(self, **kwargs):
        """Tool: Returns extracted text from the source document (Docx) or Image description."""
        content = self.context.current_context_text or ""
        
        # Integrate Canvas References
        ref_content = current_engine.get_reference_context()
        if ref_content:
            content += "\n\n[Attached Reference Documents (Canvas)]:\n" + ref_content
            
        if not content:
            return "No content available (No source loaded and no Canvas references found)."
            
        # Increased limit for multi-doc contexts (Modern LLMs handle 100k+)
        # Return first 50000 chars to be safe, but typically provide enough.
        return f"Source Content (Truncated if >50k):\n{content[:50000]}..." if len(content) > 50000 else content

    def copy_image_to_word(self, location, width_inches=2.0, **kwargs):
        """Tool: Inserts the uploaded source image into the Word document at the specified location."""
        # 防御性检查：确保有保存路径
        self._ensure_save_path()
        
        target_doc = current_engine.staging_doc if current_engine.staging_doc else current_engine.doc
        if not target_doc:
            return "Error: No Word file loaded (Canvas)."
        if not self.context.current_image_path:
            return "Error: No source image found. Please upload an image first."

        from docx.shared import Inches
        import os

        if not os.path.exists(self.context.current_image_path):
             return f"Error: Image file not found at {self.context.current_image_path}"

        try:
            width = float(width_inches)
            
            if "Table" in location:
                # Parse "Table X, Row Y, Cell Z"
                parts = location.replace("Table ", "").replace("Row ", "").replace("Cell ", "").split(",")
                t_idx, r_idx, c_idx = int(parts[0]), int(parts[1]), int(parts[2])
                
                if t_idx >= len(target_doc.tables): return "Error: Table index out of range."
                table = target_doc.tables[t_idx]
                if r_idx >= len(table.rows): return "Error: Row index out of range."
                row = table.rows[r_idx]
                
                target_cell_idx = c_idx
                # Adjust for relative logic if needed, but for image usually replace or append?
                # Let's append to the cell.
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(self.context.current_image_path, width=Inches(width))
                    
                    if current_engine.original_path:
                        target_doc.save(current_engine.original_path)
                    
                    # Signal Update
                    if hasattr(self.context, 'events'):
                        self.context.events.append({"type": "CANVAS_UPDATE"})

                    return f"Success: Image inserted into Table {t_idx}, Row {r_idx}, Cell {c_idx}"
                else:
                    return "Error: Cell index out of range."

            elif "Paragraph" in location:
                p_idx = int(location.split(" ")[1])
                if p_idx >= len(target_doc.paragraphs): return "Error: Para index out of range."
                
                para = target_doc.paragraphs[p_idx]
                run = para.add_run()
                run.add_picture(self.context.current_image_path, width=Inches(width))
                
                if current_engine.original_path:
                    target_doc.save(current_engine.original_path)

                # Signal Update
                if hasattr(self.context, 'events'):
                    self.context.events.append({"type": "CANVAS_UPDATE"})

                return f"Success: Image inserted after Paragraph {p_idx}"

            else:
                return "Error: Unsupported location format. Use 'Table X, Row Y, Cell Z' or 'Paragraph X'."

        except Exception as e:
            logger.error(f"Error copying image: {e}")
            return f"Error copying image: {str(e)}"

    def execute_document_script(self, script_code=None, **kwargs):
        """
        Tool: Executes a Python script to manipulate the document.
        The script has access to 'doc' (the python-docx Document object) and 'pd' (pandas).
        
        Usage:
        script_code should be valid python code.
        Example:
        for para in doc.paragraphs:
            if 'Start' in para.text:
                para.text = para.text.replace('Start', 'End')
        """
        # DEBUG ARGUMENTS
        log_debug(f"Tool Args Received: {list(kwargs.keys())}")
        if script_code:
             log_debug("script_code arg detected.")
        else:
             log_debug("script_code arg is None/Empty.")

        if script_code is None:
            # Try aliases
            script_code = kwargs.get("script_code") or kwargs.get("code") or kwargs.get("script")
            if not script_code and "raw_input" in kwargs:
                # Fallback: Try to parse from raw_input string if it looks like JSON
                import json
                import ast
                
                raw = kwargs["raw_input"]
                log_debug(f"Raw Input Length: {len(raw)}, Preview: {raw[:80]}...")
                
                # Strategy 1: Standard JSON parsing
                try:
                    data = json.loads(raw)
                    script_code = data.get("script_code") or data.get("code")
                    if script_code:
                        log_debug("✓ JSON parsing succeeded")
                except json.JSONDecodeError as e:
                    log_debug(f"JSON parsing failed: {e}")
                    
                    # Strategy 2: Handle common JSON issues (e.g., single quotes instead of double quotes)
                    try:
                        # Try replacing single quotes with double quotes (risky but common LLM mistake)
                        # Only do this if the string starts with { and looks like JSON
                        if raw.strip().startswith('{'):
                            # Simple heuristic: if more single quotes than double quotes at key positions, try swap
                            fixed_raw = raw.replace("'", '"')  # Risky but might work for simple cases
                            data = json.loads(fixed_raw)
                            script_code = data.get("script_code") or data.get("code")
                            if script_code:
                                log_debug("✓ JSON parsing succeeded after quote fix")
                    except:
                        pass
                        
                # Strategy 3: Regex extraction (last resort for malformed JSON)
                if not script_code:
                    log_debug("Attempting regex extraction...")
                    import re
                    
                    # IMPROVED: Use more flexible regex that can handle very long strings
                    # Match from "script_code": " up to the LAST occurrence of " before }
                    # This handles embedded quotes better
                    patterns = [
                        # Pattern 1: Standard double quotes with proper escaping
                        r'["\'](?:script_code|code)["\']\s*:\s*["\'](.+?)["\'](?:\s*[,}])',
                        # Pattern 2: Greedy match up to closing brace (for very long content)
                        r'["\'](?:script_code|code)["\']\s*:\s*["\'](.+)["\'](?=\s*})',
                        # Pattern 3: Match everything from key to end of string
                        r'(?:script_code|code)["\']\s*:\s*["\'](.+)',
                    ]
                    
                    for pattern in patterns:
                        try:
                            match = re.search(pattern, raw, re.DOTALL)
                            if match:
                                captured = match.group(1)
                                log_debug(f"Regex matched (pattern {patterns.index(pattern)+1}), captured length: {len(captured)}")
                                
                                # Try to unescape the captured content safely
                                try:
                                    # Method 1: ast.literal_eval (safest for Python string literals)
                                    script_code = ast.literal_eval(f'"{captured}"')
                                    log_debug("✓ ast.literal_eval succeeded")
                                    break
                                except Exception as ast_err:
                                    log_debug(f"ast.literal_eval failed: {ast_err}")
                                    
                                    # Method 2: Manual escape replacement (preserves UTF-8)
                                    try:
                                        script_code = (captured
                                            .replace('\\n', '\n')
                                            .replace('\\t', '\t')
                                            .replace('\\r', '\r')
                                            .replace('\\"', '"')
                                            .replace("\\'", "'")
                                            .replace('\\\\', '\\'))
                                        log_debug("✓ Manual escape replacement succeeded")
                                        break
                                    except Exception as manual_err:
                                        log_debug(f"Manual replacement failed: {manual_err}")
                                        # Use as-is as last resort
                                        script_code = captured
                                        log_debug("Using captured string as-is")
                                        break
                        except Exception as regex_err:
                            log_debug(f"Regex pattern {patterns.index(pattern)+1} failed: {regex_err}")
                            continue
                
                # Strategy 4: If all else fails and input doesn't look like JSON, use raw input directly
                if not script_code and not raw.strip().startswith('{'):
                    script_code = raw
                    log_debug("Using raw_input directly (no JSON detected)")

        if not script_code:
            log_debug("CRITICAL: No script_code found in args.")
            return "Error: No script_code provided. Please ensure you pass 'script_code' as a string containing the python code."

        if not current_engine.doc:
            return "Error: No document loaded."
        
        # 防御性检查：确保有保存路径
        self._ensure_save_path()
            
        target_doc = current_engine.staging_doc if current_engine.staging_doc else current_engine.doc
            
        try:
            log_debug(f"--- START SCRIPT EXECUTION ---")
            log_debug(f"Engine ID (Tool): {id(current_engine)}")
            log_debug(f"Script Code Preview: {script_code[:100]}...")
            
            # Sandbox / Context
            local_scope = {
                "doc": target_doc,
                "pd": pd,
                "context_df": self.context.current_df,
                "context_text": self.context.current_context_text,
                "print": lambda *args: output_buffer.append(" ".join(map(str, args))) 
            }
            
            output_buffer = []
            
            # MONKEY PATCH: Enable insert_paragraph_after for Agents
            # This logic is Verified by repro_monkeypatch.py
            from docx.text.paragraph import Paragraph
            def insert_paragraph_after(self, text=None, style=None):
                """Inserts a paragraph after this paragraph."""
                new_p = self._parent.add_paragraph(text, style)
                # Move it from end to after self
                self._parent._element.remove(new_p._element)
                self._element.addnext(new_p._element)
                return new_p
            Paragraph.insert_paragraph_after = insert_paragraph_after

            # Execute
            logger.info(f"Executing Script Code:\n{script_code}")
            exec(script_code, {}, local_scope)
            
            # Auto-save & Reload Strategy
            if current_engine.original_path:
                try:
                    # FORCE SAVE: Even if script didn't explicitly save, we save the modified object.
                    target_doc.save(current_engine.original_path)
                    logger.info(f"Script executed and saved to: {current_engine.original_path}")
                    log_debug(f"Saved to Original Path: {current_engine.original_path}")
                    
                    # RELOAD to ensure consistency (Fixes preview sync issues)
                    
                    # RELOAD to ensure consistency (Fixes preview sync issues)
                    # We MUST reload from disk to ensure the in-memory object matches the file state.
                    from docx import Document as LoadDocument
                    with open(current_engine.original_path, "rb") as f:
                        new_doc = LoadDocument(f)
                        
                    if current_engine.staging_doc:
                        current_engine.staging_doc = new_doc
                    else:
                        current_engine.doc = new_doc
                        
                    logger.info("Document reloaded from disk to sync state.")
                    log_debug("Document Reloaded from Disk.")
                    
                    # Signal Frontend to Refresh
                    if hasattr(self.context, 'events'):
                        self.context.events.append({"type": "CANVAS_UPDATE"})

                except PermissionError:
                    msg = f"Error: Could not save to '{os.path.basename(current_engine.original_path)}'. The file is likely open in Word or WPS. Please CLOSE the file and try again."
                    logger.error(msg)
                    return msg
                except Exception as e:
                    msg = f"Error saving/reloading document: {str(e)}"
                    logger.error(msg)
                    return msg
                
                return f"Script executed successfully. [Debug Path: {current_engine.original_path}] Document saved and reloaded.\nOutput: " + "\n".join(output_buffer)
                return f"Script executed successfully. [Debug Path: {current_engine.original_path}] Document saved and reloaded.\nOutput: " + "\n".join(output_buffer)
            else:
                # FALLBACK: If no original_path, save to default cache to FORCE sync and persistence!
                try:
                    # Save to backend/last_canvas.docx (User requested backend root)
                    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                    # cache_dir = os.path.join(base_dir, ".cache") # Removed subfolder per preference
                    # os.makedirs(cache_dir, exist_ok=True)
                    default_path = os.path.join(base_dir, "last_canvas.docx")
                    
                    target_doc.save(default_path)
                    current_engine.original_path = default_path # Promote to file-backed
                    logger.info(f"Memory-only doc promoted to file-backed: {default_path}")
                    log_debug(f"Fallback Save Success: {default_path}")
                    
                    # RELOAD logic (Copy-Paste from above to ensure identical behavior)
                    
                    # RELOAD logic (Copy-Paste from above to ensure identical behavior)
                    from docx import Document as LoadDocument
                    with open(default_path, "rb") as f:
                        new_doc = LoadDocument(f)
                        
                    if current_engine.staging_doc:
                        current_engine.staging_doc = new_doc
                    else:
                        current_engine.doc = new_doc
                        
                    # Signal Frontend
                    if hasattr(self.context, 'events'):
                        self.context.events.append({"type": "CANVAS_UPDATE"})
                        
                    return f"Script executed. Document saved to cache ({default_path}) and reloaded.\nOutput: " + "\n".join(output_buffer)
                except Exception as e:
                    logger.error(f"Fallback Save Failed: {e}")
                    return f"Script executed (Memory Only - Save Failed: {e}). Output: " + "\n".join(output_buffer)
                
        except Exception as e:
            return f"Script Execution Error: {str(e)}"
