import pandas as pd
import logging
import os
from werkzeug.utils import secure_filename
from core.services import current_engine
from core.services import current_engine
from .agent_engine import AgentEngine
from .prompts import PLANNER_SYSTEM_PROMPT
import json
import re

logger = logging.getLogger(__name__)

class SmartFillerService:
    def __init__(self):
        self.current_df = None
        self.current_context_text = None
        self.current_image_path = None
        self.excel_filename = None
        
        # Initial check
        self.ensure_data_loaded()
        
        # Side-channel events (e.g. for frontend refresh)
        self.events = []

    def ensure_data_loaded(self):
        """
        Global Sync: Syncs state from current_engine.reference_docs.
        Deprecates local cache loading.
        """
        # 1. Sync Excel Dataframe
        # Find the search for the FIRST excel file in references to use as primary data source
        self.current_df = None
        self.excel_filename = None
        
        # Access global engine directly
        # Note: DocxEngine stores refs as dict: {'filename':..., 'type': 'excel', 'df': ...}
        if hasattr(current_engine, 'reference_docs'):
            for ref in current_engine.reference_docs:
                if ref.get('type') == 'excel' and 'df' in ref:
                    self.current_df = ref['df']
                    self.excel_filename = ref['filename']
                    # logger.info(f"SmartFiller synced Excel from Global Engine: {ref['filename']}")
                    break
        
        # 2. Sync Context (Virtual)
        # We don't need to store text context anymore, tools will fetch it alive.
        pass

    def get_current_status(self):
        """
        Returns the current source data status (synced from global engine).
        """
        self.ensure_data_loaded()
        
        status = {
            "source_loaded": False,
            "filename": None,
            "type": None
        }
        
        if self.current_df is not None:
            status["source_loaded"] = True
            status["filename"] = self.excel_filename
            status["type"] = "excel"
            # Optional: Add preview if needed for frontend
            # status["preview"] = ...
        elif hasattr(current_engine, 'reference_docs') and len(current_engine.reference_docs) > 0:
            # Fallback for Docx refs
            ref = current_engine.reference_docs[0]
            status["source_loaded"] = True
            status["filename"] = ref['filename']
            status["type"] = ref.get('type', 'docx')
            
        return status

    def _process_raw_df(self, df_raw, filename):
        """
        Refactored logic to process raw dataframe (find header, etc)
        """
        self.excel_filename = filename
        
        # 2. Heuristic: Find the header row
        # We look for the first row that has a significant number of non-empty strings
        header_index = 0
        max_filled_cols = 0
        
        # Scan top 10 rows
        scan_limit = min(10, len(df_raw))
        for i in range(scan_limit):
            row = df_raw.iloc[i]
            valid_strings = 0
            for val in row:
                if pd.notna(val) and str(val).strip() != "":
                    valid_strings += 1
            
            # Update candidate. 
            if valid_strings > max_filled_cols:
                max_filled_cols = valid_strings
                header_index = i
        
        # 3. Check for Multi-row Header (Complex Format)
        final_header = []
        next_row_index = header_index + 1
        
        is_multi_row = False
        if next_row_index < len(df_raw):
            row_next = df_raw.iloc[next_row_index]
            # Check density of next row
            next_valid = 0
            for val in row_next:
                if pd.notna(val) and str(val).strip() != "":
                    next_valid += 1
            
            # If next row is also dense (e.g. > 50% of max dense), assume multi-row
            if next_valid > (max_filled_cols * 0.5):
                is_multi_row = True
        
        if is_multi_row:
            # Merge Row H and Row H+1
            row0 = df_raw.iloc[header_index].tolist()
            row1 = df_raw.iloc[next_row_index].tolist()
            
            # Forward fill row0 (handle merged cells horizontal)
            last_val = ""
            filled_row0 = []
            for val in row0:
                val_str = str(val).strip() if pd.notna(val) else ""
                if val_str == "":
                    filled_row0.append(last_val)
                else:
                    filled_row0.append(val_str)
                    last_val = val_str
            
            merged_header = []
            for i in range(len(filled_row0)):
                h1 = filled_row0[i]
                h2 = str(row1[i]).strip() if pd.notna(row1[i]) else ""
                
                # Combine: "Recruitment Condition - Age" or just "Age" if H1 is same/empty
                if h1 and h2 and h1 != h2:
                    merged_header.append(f"{h1} {h2}") # Space separator is safer than hyphen
                elif h2:
                    merged_header.append(h2)
                elif h1:
                    merged_header.append(h1)
                else:
                    merged_header.append(f"Unnamed_{i}")
            
            final_header = merged_header
            # Slice data start from row after H+1
            df = df_raw[next_row_index + 1:].copy()
        else:
            # Layout: Single Header Row
            if header_index >= 0:
                row = df_raw.iloc[header_index]
                final_header = [str(val).strip() if pd.notna(val) else f"Unnamed_{j}" for j, val in enumerate(row)]
                df = df_raw[header_index + 1:].copy()
            else:
                 # Fallback
                final_header = [f"Col_{j}" for j in range(len(df_raw.columns))]
                df = df_raw.copy()

        df.columns = final_header
        df = df.fillna("").infer_objects(copy=False)
        df = df.reset_index(drop=True)

        # Store for later use
        self.current_df = df
        
        logger.info(f"Excel parsed. Filename: {filename}, Detected Header Row: {header_index}, Multi-row: {is_multi_row}, Columns: {df.columns.tolist()}")

        return {
            "filename": filename,
            "columns": df.columns.tolist(),
            "total_rows": len(df),
            "preview": df.head(5).to_dict(orient='records')
        }

    def parse_source(self, file_storage):
        """
        DEPRECATED: Use Global Upload (/canvas/upload_reference).
        This method is kept only to prevent API 500s if old UI is used, but it does nothing?
        Actually, we should just disable it or redirect logic to global ADD if needed.
        But since UI is removing it, we just warn.
        """
        logger.warning("Deprecated parse_source called. Please use Global Upload.")
        return {"message": "Deprecated. Use top-right upload button."}

    def parse_excel(self, file_storage):
        """Wrapper for backward compatibility"""
        return self.parse_source(file_storage)

    def _process_docx(self, filepath, filename):
        """Extracts text and images from Docx"""
        try:
            import docx
            import zipfile
            import shutil
            
            # 1. Extract Text
            doc = docx.Document(filepath)
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Also extract tables text
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            # 2. Extract Images (Multi-modal)
            image_note = ""
            try:
                # Prepare extraction dir
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                extract_dir = os.path.join(base_dir, ".cache", "extracted_images")
                if os.path.exists(extract_dir):
                    shutil.rmtree(extract_dir) # Clear old images
                os.makedirs(extract_dir, exist_ok=True)

                extracted_count = 0
                first_image_path = None
                
                with zipfile.ZipFile(filepath) as z:
                    for name in z.namelist():
                        if name.startswith("word/media/") and name.lower().endswith(('.png', '.jpg', '.jpeg')):
                            # Extract
                            target_name = os.path.basename(name)
                            target_path = os.path.join(extract_dir, target_name)
                            with open(target_path, "wb") as f:
                                f.write(z.read(name))
                            
                            extracted_count += 1
                            if not first_image_path:
                                first_image_path = target_path
                
                if extracted_count > 0:
                    self.current_image_path = first_image_path
                    image_note = f"\n\n[System Info: Detected {extracted_count} images in document. The first image is available for insertion. Use 'copy_image_to_word' tool.]"
                    logger.info(f"Extracted {extracted_count} images. Primary: {first_image_path}")
                else:
                    self.current_image_path = None

            except Exception as img_e:
                logger.warning(f"Image extraction failed: {img_e}")

            # Store as "Context"
            new_content = f"\n\n--- [Document: {filename}] ---\n{full_text}{image_note}"
            
            if self.current_context_text:
                self.current_context_text += new_content
            else:
                self.current_context_text = new_content
            # Keep existing DF if any
            # self.current_df = None 
            
            logger.info(f"Parsed Docx: {filename}, Length: {len(full_text)}, Images: {self.current_image_path is not None}")
            return {
                "filename": filename,
                "type": "docx",
                "preview": [{"content": line} for line in text_content[:5]],
                "total_lines": len(text_content),
                "has_images": self.current_image_path is not None
            }
        except Exception as e:
            logger.error(f"Docx parsing failed: {e}")
            raise e

    def _process_image(self, filepath, filename):
        """Handles Image upload. Sets up path for copying. OCR placeholder."""
        logger.info(f"Image uploaded: {filename}")
        self.current_image_path = filepath
        self.current_df = None
        
        # In MVP, we don't have local OCR. 
        # We set this text so the Agent knows it has an image but cannot read it yet.
        self.current_context_text = f"[Image File: {filename}]\\n(Image content extraction/OCR is not enabled locally. You can use 'copy_image_to_word' to insert this image into the document.)"
        
        return {
            "filename": filename,
            "type": "image",
            "preview": [{"content": "Image uploaded. Ready to copy to Word."}],
            "path": filepath
        }


    def get_recommendations(self):
        """
        Analyzes the current document in DocxEngine and recommends mappings 
        based on the loaded Excel dataframe.
        """
        # Ensure data is loaded
        self.ensure_data_loaded()

        if self.current_df is None:
            return {"error": "No Excel file loaded"}
            
        if not current_engine.doc:
            return {"error": "No document loaded in canvas"}

        excel_columns = self.current_df.columns.tolist()
        recommendations = {
            "tables": [],
            "placeholders": []
        }

        # 1. Match Tables
        # Identify tables in Docx that have headers matching Excel columns
        for i, table in enumerate(current_engine.doc.tables):
            # rudimentary check: look at first row
            try:
                if not table.rows: continue
                first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                # Calculate intersection
                common = set(excel_columns) & set(first_row_cells)
                if len(common) > 0:
                    recommendations['tables'].append({
                        "table_index": i,
                        "matches": list(common),
                        "confidence": len(common) / len(first_row_cells) if len(first_row_cells) > 0 else 0
                    })
            except Exception as e:
                logger.warning(f"Error scanning table {i}: {e}")
                continue

        return recommendations

    def fill_table(self, table_index):
        """
        Fills the specified Word table (by index) with data from the loaded Excel dataframe.
        """
        self._log_debug(f"--- START FILL TABLE (Index: {table_index}) ---")
        
        # Ensure data is loaded
        self.ensure_data_loaded()
        
        if self.current_df is None:
            self._log_debug("Error: No Excel file loaded.")
            raise ValueError("No Excel file loaded. Please upload one first.")
        
        if not current_engine.doc:
            self._log_debug("Error: No document loaded.")
            raise ValueError("No document loaded in canvas.")

        try:
            table = current_engine.doc.tables[table_index]
        except IndexError:
            self._log_debug(f"Error: Table index {table_index} not found.")
            raise ValueError(f"Table index {table_index} not found.")

        # Mapping: Word Column Index -> Excel Column Name
        # We derive this from the first row of the table
        if not table.rows:
             self._log_debug("Error: Target table has no rows.")
             raise ValueError("Target table has no rows (cannot map headers).")
        
        header_row = table.rows[0]
        col_map = {} # { word_col_idx: excel_col_name }
        
        # Build map based on name match
        headers_found = []
        for idx, cell in enumerate(header_row.cells):
            text = cell.text.strip()
            if text in self.current_df.columns:
                col_map[idx] = text
                headers_found.append(text)
        
        self._log_debug(f"Mapped Columns: {headers_found}")

        if not col_map:
            self._log_debug("Error: No matching headers found.")
            raise ValueError("No matching headers found between Word table and Excel columns.")

        # Data Injection
        rows_added = 0
        for _, row_data in self.current_df.iterrows():
            # Add a NEW row to the table
            new_row = table.add_row() 
            rows_added += 1
            # Fill cells based on map
            for col_idx, excel_col in col_map.items():
                if col_idx < len(new_row.cells):
                    val = str(row_data[excel_col])
                    new_row.cells[col_idx].text = val
        
        self._log_debug(f"Filled {rows_added} rows.")
        
        # Save the modified document
        if current_engine.original_path:
             current_engine.doc.save(current_engine.original_path)
             self._log_debug(f"Document saved to: {current_engine.original_path}")
        else:
             logger.warning("No original path to save to (in-memory document?)")
             self._log_debug("Warning: No original path to save to.")
             
        return {"status": "success", "rows_added": len(self.current_df)}
    # --- Tool Use Methods ---

    def read_excel_summary(self):
        """Tool: Returns summary of loaded Excel."""
        self.ensure_data_loaded()
        if self.current_df is None:
            return "No Excel file loaded."
        columns = ", ".join(self.current_df.columns.tolist())
        rows = len(self.current_df)
        return f"Excel loaded. Columns: [{columns}]. Total Rows: {rows}. Top 3 rows data: {self.current_df.head(3).to_dict(orient='records')}"

    def find_anchor_in_word(self, text):
        """Tool: Finds the location of text in Word. Returns 'Table X, Row Y, Cell Z' or 'Paragraph X'."""
        if not current_engine.doc:
            return "No document loaded."
        
        # 1. Search Tables
        for t_idx, table in enumerate(current_engine.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if text in cell.text:
                        return f"Table {t_idx}, Row {r_idx}, Cell {c_idx}"
        
        # 2. Search Paragraphs
        for p_idx, para in enumerate(current_engine.doc.paragraphs):
            if text in para.text:
                return f"Paragraph {p_idx}"
                
        return f"Text '{text}' not found."

    def write_word_content(self, location, text, relative="right"):
        """Tool: Writes text to a location. Relative='right' means next cell."""
        if not current_engine.doc:
            return "No document loaded."

        try:
            if "Table" in location:
                # Parse "Table 0, Row 0, Cell 0"
                parts = location.replace("Table ", "").replace("Row ", "").replace("Cell ", "").split(",")
                t_idx, r_idx, c_idx = int(parts[0]), int(parts[1]), int(parts[2])
                
                table = current_engine.doc.tables[t_idx]
                row = table.rows[r_idx]
                
                target_cell_idx = c_idx
                if relative == "right":
                    target_cell_idx += 1
                
                if target_cell_idx < len(row.cells):
                    row.cells[target_cell_idx].text = str(text)
                    if current_engine.original_path:
                        current_engine.doc.save(current_engine.original_path)
                    return f"Written '{text}' to Table {t_idx}, Row {r_idx}, Cell {target_cell_idx}"
                else:
                    return "Target cell out of range."
                    
            elif "Paragraph" in location:
                # Append to paragraph? Or replace? 
                # Simplest for MVP: Append
                p_idx = int(location.split(" ")[1])
                para = current_engine.doc.paragraphs[p_idx]
                para.add_run(f" {text}")
                if current_engine.original_path:
                    current_engine.doc.save(current_engine.original_path)
                return f"Appended '{text}' to Paragraph {p_idx}"
                
        except Exception as e:
            return f"Error writing content: {e}"

        return "Invalid location format."

    def _log_debug(self, message):
        """Helper to append to debug log file"""
        try:
            log_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "smart_filler_debug.log")
            
            # Ensure the directory exists (should be backend/)
            # Just incase
            os.makedirs(os.path.dirname(log_path), exist_ok=True)
            
            with open(log_path, "a", encoding="utf-8") as f:
                import datetime
                timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                f.write(f"[{timestamp}] {message}\\n")
                f.flush() # Force write
        except Exception as e:
            # Fallback print if file logging fails
            print(f"[Fallback Log] {message} (Error: {e})")

    def run_agent_task(self, instruction, model_config=None, plan=None):
        """
        Agent Loop: Delegated to AgentEngine ("Old Yang").
        """
        self.ensure_data_loaded() # Ensure Loaded
        self.events = [] # Clear side-channel events for this run
        
        try:
            engine = AgentEngine(self)
            result = engine.run(instruction, model_config, plan)
            
            # Attach accumulated events (e.g. CANVAS_UPDATE)
            if isinstance(result, dict):
                result['events'] = list(self.events) # copy
                
            return result
        except Exception as e:
            logger.error(f"Agent Engine Failed: {e}")
            self._log_debug(f"Agent Engine Failed: {e}")
            return {"status": "error", "message": str(e), "events": []}

    def generate_plan(self, instruction, model_config=None):
        """
        Generates a structured plan for the given instruction using the Planner Agent.
        """
        self.ensure_data_loaded()
        
        try:
            engine = AgentEngine(self)
            plan_json = engine.generate_plan_only(instruction, model_config)
            
            # --- FORCE FINAL SYNC STEP ---
            if isinstance(plan_json, list):
                # Check if last step is already a save/sync (heuristic)
                last_step_desc = plan_json[-1].get('description', '') if plan_json else ''
                if "刷新" not in last_step_desc and "Save" not in last_step_desc:
                    next_step_num = len(plan_json) + 1
                    plan_json.append({
                        "step": next_step_num,
                        "description": "强制同步：执行空脚本以确保所有更改已保存并触发画布刷新。",
                        "tool_hint": "execute_document_script"
                    })
            # -----------------------------
            
            return {"status": "success", "plan": plan_json}
            
        except Exception as e:
            logger.error(f"Planning Failed: {e}")
            return {"status": "error", "message": str(e)}

smart_filler_service = SmartFillerService()
