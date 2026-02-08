"""
Canvas Converter: Tiptap JSON ↔ DOCX 双向转换
为快速画布和专业画布提供无缝数据互通
"""
import io
import logging
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def tiptap_to_docx(json_content: Dict[str, Any]) -> io.BytesIO:
    """
    将Tiptap JSON转换为DOCX文件流
    
    支持的节点类型:
    - paragraph (段落)
    - heading (标题 level 1-6)
    - bulletList / listItem (无序列表)
    - orderedList / listItem (有序列表)
    - hardBreak (换行)
    
    支持的标记:
    - bold (粗体)
    - italic (斜体)
    - underline (下划线)
    - strike (删除线)
    """
    doc = Document()
    
    # 处理content数组
    content = json_content.get('content', [])
    
    for block in content:
        block_type = block.get('type')
        
        if block_type == 'paragraph':
            _process_paragraph(doc, block)
        
        elif block_type == 'heading':
            level = block.get('attrs', {}).get('level', 1)
            text_content = _extract_text(block)
            doc.add_heading(text_content, level=min(level, 3))
        
        elif block_type == 'bulletList':
            _process_list(doc, block, ordered=False)
        
        elif block_type == 'orderedList':
            _process_list(doc, block, ordered=True)
    
    # 返回BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _process_paragraph(doc: Document, block: Dict):
    """处理段落节点，应用文本标记"""
    para = doc.add_paragraph()
    content = block.get('content', [])
    
    for item in content:
        if item.get('type') == 'text':
            text = item.get('text', '')
            run = para.add_run(text)
            
            # 应用标记
            marks = item.get('marks', [])
            for mark in marks:
                mark_type = mark.get('type')
                if mark_type == 'bold':
                    run.bold = True
                elif mark_type == 'italic':
                    run.italic = True
                elif mark_type == 'underline':
                    run.underline = True
                elif mark_type == 'strike':
                    run.font.strike = True
        
        elif item.get('type') == 'hardBreak':
            para.add_run('\n')

def _process_list(doc: Document, block: Dict, ordered: bool = False):
    """处理列表节点"""
    content = block.get('content', [])
    
    for idx, item in enumerate(content):
        if item.get('type') == 'listItem':
            text_content = _extract_text(item)
            prefix = f"{idx + 1}. " if ordered else "• "
            doc.add_paragraph(prefix + text_content)

def _extract_text(block: Dict) -> str:
    """递归提取文本内容"""
    content = block.get('content', [])
    text_parts = []
    
    for item in content:
        if item.get('type') == 'text':
            text_parts.append(item.get('text', ''))
        elif 'content' in item:
            text_parts.append(_extract_text(item))
    
    return ''.join(text_parts)

# ==================== DOCX → Tiptap JSON ====================

def docx_to_tiptap(file_stream: io.BytesIO) -> Dict[str, Any]:
    """
    将DOCX转换为Tiptap JSON格式
    
    返回格式:
    {
        "type": "doc",
        "content": [...]
    }
    """
    document = Document(file_stream)
    content = []
    
    # Helper to iterate elements in order
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    def iter_block_items(parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    count_para = 0
    count_table = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            count_para += 1
            # Check style mainly for headings
            if block.style.name.startswith('Heading'):
                try:
                    level = int(block.style.name[-1])
                    content.append({
                        "type": "heading",
                        "attrs": {"level": level},
                        "content": [{"type": "text", "text": block.text}]
                    })
                except ValueError:
                    content.append(_para_to_tiptap(block))
            else:
                content.append(_para_to_tiptap(block))
        
        elif isinstance(block, Table):
            count_table += 1
            # Simple table extraction: convert rows to paragraphs
            # TODO: Future support for actual Tiptap tables
            for row in block.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                
                if row_text:
                    # Join cell text with separator
                    line = " | ".join(row_text)
                    content.append({
                        "type": "paragraph",
                        "content": [{"type": "text", "text": line}]
                    })
    
    print(f"[DEBUG] DOCX Import: {count_para} paragraphs, {count_table} tables parsed.")
    
    # Fallback: If content is still empty (e.g. text boxes, weird formatting), extract raw text
    if not content:
        print("[DEBUG] Content empty after structural parse. Attempting raw text fallback.")
        full_text = []
        for element in document.element.body.iter():
            if element.tag.endswith('t'):
                if element.text:
                    full_text.append(element.text)
        
        if full_text:
            joined_text = "".join(full_text)
            # Create a single paragraph with all text
            content.append({
                "type": "paragraph",
                "content": [{"type": "text", "text": joined_text}]
            })
            print(f"[DEBUG] Fallback: Extracted {len(joined_text)} chars of raw text.")

    print(f"[DEBUG] Final Tiptap Content Nodes: {len(content)}")

    return {
        "type": "doc",
        "content": content
    }

def _para_to_tiptap(para) -> Dict:
    """将DOCX段落转换为Tiptap段落节点"""
    text_content = []
    
    for run in para.runs:
        # Filter out empty text runs which cause ProseMirror errors
        if not run.text:
            continue
            
        text_node = {"type": "text", "text": run.text}
        
        # 提取marks
        marks = []
        if run.bold:
            marks.append({"type": "bold"})
        if run.italic:
            marks.append({"type": "italic"})
        if run.underline:
            marks.append({"type": "underline"})
        if run.font.strike:
            marks.append({"type": "strike"})
        
        if marks:
            text_node["marks"] = marks
        
        text_content.append(text_node)
    
    # If paragraph is empty, return it without content (ProseMirror handles this as empty block)
    if not text_content:
        return {"type": "paragraph"}
        
    return {
        "type": "paragraph",
        "content": text_content
    }

# ==================== Smart Gov Export ====================

import re
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

def tiptap_to_smart_docx(json_content: Dict[str, Any]) -> io.BytesIO:
    """
    智能公文格式导出
    规则：
    1. 一、   -> 黑体，三号 (通常)，一级标题
    2. （一） -> 方正楷体_GBK，三号 (通常)，二级标题
    3. 1      -> 仿宋_GB2312，三号 (通常)，三级标题
    4. 正文   -> 方正仿宋_GBK，三号，28磅固定行距
    """
    doc = Document()
    
    # 基础样式设定 (正文)
    style = doc.styles['Normal']
    style.font.name = '方正仿宋_GBK'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(28)
    
    content = json_content.get('content', [])
    
    # 正则表达式 (兼容更多格式，如全角空格)
    re_h1 = re.compile(r"^\s*[一二三四五六七八九十]+、")
    re_h2 = re.compile(r"^\s*[（(][一二三四五六七八九十]+[)）]")
    re_h3 = re.compile(r"^\s*[0-9]+[.\u3002、 ]") # 1. 1。 1、 1 
    
    for block in content:
        # 提取纯文本用于判断
        full_text = _extract_text(block).strip()
        if not full_text:
            continue
            
        block_type = block.get('type')
        
        # 默认作为正文处理
        p = doc.add_paragraph()
        
        # 智能层级判断
        if re_h1.match(full_text):
            # 一级标题：黑体，三号
            _set_run_style(p.add_run(full_text), '黑体', 16, bold=True)
        elif re_h2.match(full_text):
            # 二级标题：方正楷体_GBK，三号
            _set_run_style(p.add_run(full_text), '方正楷体_GBK', 16, bold=True)
        elif re_h3.match(full_text):
            # 三级标题：方正仿宋_GBK，三号
            _set_run_style(p.add_run(full_text), '方正仿宋_GBK', 16, bold=True)
        else:
            # 正文内容：方正仿宋_GBK，三号
            _process_paragraph_smart(p, block)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _process_paragraph_smart(paragraph, block):
    """处理正文段落，保持原有加粗/斜体，但强制字体"""
    content = block.get('content', [])
    for item in content:
        if item.get('type') == 'text':
            text = item.get('text', '')
            run = paragraph.add_run(text)
            
            # Apply base Gov Font (方正仿宋_GBK)
            _set_run_style(run, '方正仿宋_GBK', 16) 
            
            # Apply Marks (Overlay)
            marks = item.get('marks', [])
            for mark in marks:
                if mark.get('type') == 'bold': run.bold = True
                if mark.get('type') == 'italic': run.italic = True
                if mark.get('type') == 'underline': run.underline = True

def _set_run_style(run, font_name, size_pt, bold=False):
    """Helper to set Chinese font and size correctly for both ASCII and EastAsia"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    
    # 强制设置 XML 属性以确保中西文都使用该字体
    rPr = run._element.get_or_add_rPr()
    
    # 清除可能存在的旧字体设置
    # rFonts = rPr.find(qn('w:rFonts'))
    # if rFonts is not None:
    #     rPr.remove(rFonts)
        
    # 设置新的字体
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name) # Complex Script

