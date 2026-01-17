
import json
import io
import os
import re
import logging
import hashlib
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import RGBColor
from typing import List, Dict, Any, Optional
from .markdown_utils import apply_markdown_to_paragraph

# Configure logging
log_path = os.getenv("LOG_PATH", "docx_engine_debug.log")
if not os.path.isabs(log_path):
    pass

logging.basicConfig(filename=log_path, level=logging.INFO, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

class DocxEngine:
    def __init__(self):
        self.doc = None
        self.staging_doc = None # Temporary document for previews
        self.original_path = None
        self.reference_docs = [] # List of {filename: str, doc: Document}
        
        # Phase 5: DOCX Caching Infrastructure
        self._cache = {
            'file_hash': None,      # MD5 hash of loaded file
            'preview_data': None,   # Cached paragraph preview
            'toc': None,            # Cached table of contents
            'full_text': None,      # Cached full text extraction
            'reference_hashes': {}  # Track reference doc hashes
        }
        self._modified_paras = set()  # Track modified paragraphs for incremental save
        
        # Persistence: Try to reload last canvas
        self._try_reload_last_canvas()

    def _try_reload_last_canvas(self):
        try:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            cache_dir = os.path.join(base_dir, ".cache")
            cache_path = os.path.join(cache_dir, "last_canvas.docx")
            
            if os.path.exists(cache_path):
                logging.info(f"Docs Persistence: Reloading {cache_path}")
                with open(cache_path, "rb") as f:
                    self.doc = Document(f)
                self.original_path = cache_path # Or restore original name if stored?
                logging.info("Docs Persistence: Canvas reloaded.")
        except Exception as e:
            logging.warning(f"Docs Persistence Failed: {e}")

    def load_document(self, file_stream):
        """Load document with MD5-based caching to avoid redundant parsing"""
        # Read content for MD5 calculation
        content = file_stream.read()
        file_hash = hashlib.md5(content).hexdigest()
        
        # Check cache
        if self._cache['file_hash'] == file_hash and self.doc is not None:
            logging.info(f"[Cache HIT] Skipping DOCX parse (hash: {file_hash[:8]}...)")
            # Clear references on new document load to sync with frontend
            self.reference_docs = []
            return self.doc
        
        # Cache miss - parse document
        logging.info(f"[Cache MISS] Parsing DOCX (hash: {file_hash[:8]}...)")
        self.doc = Document(io.BytesIO(content))
        self.staging_doc = None
        
        # Update cache
        self._cache['file_hash'] = file_hash
        self._invalidate_cache()  # Clear derived caches
        self._modified_paras.clear()
        
        # Clear references on new document load to sync with frontend
        self.reference_docs = []
        return self.doc
    
    def _invalidate_cache(self, keys=None):
        """Invalidate specific cache keys or all derived caches"""
        if keys is None:
            # Invalidate all derived caches
            self._cache['preview_data'] = None
            self._cache['toc'] = None
            self._cache['full_text'] = None
        else:
            for key in keys:
                if key in self._cache:
                    self._cache[key] = None

    def load_from_path(self, path: str):
        """
        Loads a document from a local file path.
        """
        with open(path, "rb") as f:
            self.doc = Document(f)
        self.original_path = path
        self.staging_doc = None
        return self.doc

    def load_from_text(self, text: str, preserve_references=False):
        """
        Creates a new document from the provided text.
        
        Args:
            text: 文本内容
            preserve_references: 是否保留现有参考文档（默认False以保持向后兼容）
        """
        self.doc = Document()
        # Set default style to something reasonable for Chinese if possible,
        # but python-docx defaults are usually okay-ish. 
        # Ideally we'd set a style, but for now simple paragraphs.
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                self.doc.add_paragraph(line)
        
        self.staging_doc = None
        
        # 根据参数决定是否保留参考文档
        if not preserve_references:
            self.reference_docs = []  # 清空参考文档
        
        # ⚠️ 关键修复：确保有有效的保存路径
        # 避免 Agent 写入操作因路径为 None 而失败
        if self.original_path is None:
            import os
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            cache_dir = os.path.join(base_dir, ".cache")
            if not os.path.exists(cache_dir):
                os.makedirs(cache_dir, exist_ok=True)
            self.original_path = os.path.join(cache_dir, "canvas_from_text.docx")
            
            # 立即保存以建立文件（避免后续保存时找不到路径）
            self.doc.save(self.original_path)
        
        return self.doc

    def save_to_path(self, path: str):
        """
        Saves document with intelligent strategy:
        - Small edits (<10 paras): Incremental save (faster)
        - Large edits: Full save (safer)
        """
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc:
            return False
        
        # Determine save strategy
        mod_count = len(self._modified_paras)
        
        if mod_count > 0 and mod_count < 10:
            # Incremental save - just write to disk faster
            logging.info(f"[Incremental Save] {mod_count} paragraphs modified")
            # For now, still do full save but with optimization flag
            # Future: implement XML-level partial write
            target_doc.save(path)
        else:
            # Full save
            if mod_count >= 10:
                logging.info(f"[Full Save] {mod_count} paragraphs modified (threshold exceeded)")
            else:
                logging.info(f"[Full Save] Clean document")
            target_doc.save(path)
        
        # Update cache hash after save
        try:
            with open(path, 'rb') as f:
                content = f.read()
                new_hash = hashlib.md5(content).hexdigest()
                self._cache['file_hash'] = new_hash
                logging.info(f"[Cache] Updated file hash: {new_hash[:8]}...")
        except Exception as e:
            logging.warning(f"Failed to update cache hash: {e}")
        
        # Clear modification tracking
        self._modified_paras.clear()
        
        return True
    
    def track_modification(self, para_index: int):
        """Track paragraph modification for incremental save optimization"""
        self._modified_paras.add(para_index)
        # Invalidate affected caches
        self._invalidate_cache(['preview_data', 'full_text'])

    def add_reference_doc(self, file_stream, filename):
        """
        Loads a reference document and stores it.
        Also extracts images and converts to Markdown for better context.
        """
    def add_reference_doc(self, file_stream, filename):
        """
        Loads a reference document and stores it.
        Supports .docx (legacy) and .xlsx (new Unified Upload).
        """
        try:
            # 1. Handle Excel Files
            if filename.lower().endswith(('.xlsx', '.xls', '.csv')):
                import pandas as pd
                file_stream.seek(0)
                try:
                    if filename.lower().endswith('.csv'):
                        df = pd.read_csv(file_stream)
                    else:
                        df = pd.read_excel(file_stream)
                    
                    # Store as structured reference
                    self.reference_docs.append({
                        "filename": filename,
                        "type": "excel",
                        "df": df,
                        "markdown": f"### Excel File: {filename}\nHeaders: {list(df.columns)}\nTop 5 Rows:\n{df.head(5).to_markdown()}",
                        "doc": None
                    })
                    logging.info(f"Loaded Excel Reference: {filename}")
                    return True, "Success (Excel Loaded)"
                except Exception as e:
                    return False, f"Failed to parse Excel: {e}"

            # 2. Handle Docx Files (Existing Logic)
            # Save stream to temp file for mammoth
            import tempfile
            import shutil
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                shutil.copyfileobj(file_stream, tmp)
                tmp_path = tmp.name
            
            # Load into python-docx for structure (optional but good for backup)
            file_stream.seek(0)
            ref_doc = Document(file_stream)
            
            # Extract content with images using mammoth
            image_output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'images')
            if not os.path.exists(image_output_dir):
                os.makedirs(image_output_dir, exist_ok=True)
            
            markdown_content = self._convert_to_markdown_with_images(tmp_path, image_output_dir)
            
            # Cleanup temp file
            try:
                os.unlink(tmp_path)
            except:
                pass

            # Check for error in markdown extraction
            if markdown_content and markdown_content.startswith("Error extracting"):
                logging.warning(f"Mammoth extraction failed for {filename}: {markdown_content}. Fallback to plain text.")
                markdown_content = None
            
            self.reference_docs.append({
                "filename": filename,
                "type": "docx",
                "doc": ref_doc,
                "markdown": markdown_content
            })
            return True, "Success"
        except Exception as e:
            logging.error(f"Failed to load reference doc {filename}: {e}")
            return False, str(e)

    def reset(self):
        """
        Resets the engine to initial state.
        Also cleans up temporary images in static/images.
        """
        self.doc = None
        self.staging_doc = None
        self.original_path = None
        self.reference_docs = []
        
        # Cleanup Logic
        try:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            static_img_dir = os.path.join(base_dir, 'static', 'images')
            if os.path.exists(static_img_dir):
                logging.info(f"Cleaning up image directory: {static_img_dir}")
                for filename in os.listdir(static_img_dir):
                    file_path = os.path.join(static_img_dir, filename)
                    try:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                    except Exception as e:
                        logging.error(f"Failed to delete {file_path}. Reason: {e}")
        except Exception as e:
            logging.error(f"Error during reset cleanup: {e}")
            
        return True
        
    def get_reference_list(self) -> list:
        return [ref['filename'] for ref in self.reference_docs]
        
    def remove_reference_doc(self, filename: str) -> bool:
        initial_len = len(self.reference_docs)
        self.reference_docs = [ref for ref in self.reference_docs if ref['filename'] != filename]
        return len(self.reference_docs) < initial_len

    def get_reference_context(self) -> str:
        """
        Returns a text representation of all reference documents for the LLM.
        Prioritizes Markdown content (which includes image links).
        """
        context_parts = []
        for ref in self.reference_docs:
            context_parts.append(f"--- Reference File: {ref['filename']} ---")
            if 'markdown' in ref and ref['markdown']:
                context_parts.append(ref['markdown'])
            else:
                # Fallback to plain text
                # Fallback to plain text with Real Page Detection (Hybrid)
                char_count_since_last_page = 0
                page_num = 1
                context_parts.append(f"[第 1 页]")
                
                for para in ref['doc'].paragraphs:
                    para_text_buffer = ""
                    page_break_found = False
                    
                    # 1. Try to detect breaks in runs
                    for run in para.runs:
                        run_xml = run._element.xml
                        # Check for hard break or soft break
                        if 'w:lastRenderedPageBreak' in run_xml or 'w:br' in run_xml and 'w:type="page"' in run_xml:
                             page_num += 1
                             para_text_buffer += f"\n\n[第 {page_num} 页]\n"
                             page_break_found = True
                             char_count_since_last_page = 0
                        
                        para_text_buffer += run.text
                    
                    # Explicit Textual Page Break Detection
                    # Case A: Visual Marker "............ 分页符 ............."
                    if "分页符" in para_text_buffer and "....." in para_text_buffer:
                        if not page_break_found:
                             page_num += 1
                             para_text_buffer += f"\n\n[第 {page_num} 页]\n"
                             char_count_since_last_page = 0
                             page_break_found = True

                    # Case B: Footer/Page Number Line detection (Aggressive)
                    # Support: "57", "- 57 -", "第 57 页", "Page 57", "· 57 ·"
                    # Regex: Optional [Page/第/-/.], number, Optional [Page/页/-/.], End of line.
                    p_clean = para_text_buffer.strip()
                    page_num_match = re.search(r'^\s*(?:page|第|[·\.\-_])?\s*(\d+)\s*(?:page|页|[·\.\-_])?\s*$', p_clean, re.IGNORECASE)
                    
                    if page_num_match:
                        try:
                            detected_num = int(page_num_match.group(1))
                            # Sanity checks:
                            # 1. Page number < 2000 (Avoid years like 2023)
                            # 2. Line length < 20 chars (Avoid long sentences that happen to match weirdly, though regex anchored ^$)
                            if 0 < detected_num < 2000 and len(p_clean) < 20:
                                page_num = detected_num
                                if not page_break_found:
                                    # This implies we just crossed into this page
                                    pass
                        except:
                            pass
                    if not page_num_match:
                         # Try pure number if centered? Risk of false positive. Stick to "第 N 页".
                         pass
                    
                    if page_num_match:
                        # Found explicit page number!
                        detected_num = int(page_num_match.group(1))
                        if detected_num > 0:
                            page_num = detected_num
                            # Don't increment, just set. And this line IS the page number.
                            # We can mark it as such.
                            # If we just saw a break, maybe this confirms it.
                            if not page_break_found:
                                # If we didn't just break, maybe we should have?
                                # This indicates NEW page context usually starts here or just before.
                                pass

                        
                    # 2. Append text unique to this paragraph
                    if para_text_buffer.strip():
                        context_parts.append(para_text_buffer)
                        char_count_since_last_page += len(para_text_buffer)
                        
                        # 3. Failsafe: If > 1200 chars without a break, force an estimate
                        # Only if we haven't seen a real break recently
                        if char_count_since_last_page > 1200:
                            page_num += 1
                            context_parts.append(f"\n[第 {page_num} 页 (估算)]\n")
                            char_count_since_last_page = 0
                            
                # Append tables at the end (simplification)
                for table in ref['doc'].tables:
                    context_parts.append("\n[表格数据]")
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        context_parts.append(row_text)

            context_parts.append("--- 文档结束 ---\n")
        
        return "\n".join(context_parts)

    def get_content_by_indices(self, indices_list: List[Dict[str, Any]]) -> str:
        """
        Retrieves content from reference docs based on paragraph indices.
        Wrapper for get_content_by_toc_items.
        """
        return self.get_content_by_toc_items(indices_list)

    def get_reference_structure(self) -> List[Dict[str, Any]]:
        """
        Returns a combined TOC for all reference documents.
        """
        combined_toc = []
        
        for doc_idx, ref in enumerate(self.reference_docs):
            doc = ref['doc']
            filename = ref['filename']
            
            # Use enhanced extraction on reference doc
            toc = self._extract_toc_from_doc(doc)
            
            # Add context to items
            for item in toc:
                item['filename'] = filename
                item['doc_idx'] = doc_idx
                # Create a composite ID for LLM clarity if needed, 
                # but for now we rely on LLM returning the full object or 'doc_idx' if prompt asks?
                # Actually LLM prompt currently just asks for start/end/title.
                # We need to make sure LLM can return enough info to identify the doc.
                # We will handle this by making the "ID" unique in the prompt in llm_engine.
                combined_toc.append(item)
                
        return combined_toc

    def get_content_by_toc_items(self, items: List[Dict[str, Any]]) -> str:
        """
        Retrieves content based on selected TOC items.
        """
        final_context = []
        
        for item in items:
            doc_idx = item.get('doc_idx')
            
            # If doc_idx is missing, try to find by filename or fallback to 0
            if doc_idx is None:
                fname = item.get('filename')
                if fname:
                    for i, ref in enumerate(self.reference_docs):
                        if ref['filename'] == fname:
                            doc_idx = i
                            break
            
            # Fallback to 0 if single ref doc
            if doc_idx is None and len(self.reference_docs) == 1:
                doc_idx = 0
            
            if doc_idx is None or doc_idx >= len(self.reference_docs):
                continue
                
            ref = self.reference_docs[doc_idx]
            doc = ref['doc']
            filename = ref['filename']
            start = item['start'] # LLM might return 'start' key
            # In TOC construction we use 'id' as start.
            if 'id' in item: start = item['id']
            
            end = item.get('end', start + 20) # Default
            if 'end_id' in item: end = item['end_id']
            
            # Extract
            section_text = self._extract_range_with_pages(doc, start, end)
            final_context.append(f"--- 来源：{filename} | 章节：{item.get('title')} ---\n{section_text}\n")
            
        return "\n".join(final_context)

    def _extract_range_with_pages(self, doc, start_idx, end_idx, limit=3000):
        """
        Extracts text from paragraph range with Page Numbers.
        Includes Character-Count Fallback for docs without explicit breaks.
        """
        content = []
        current_page = 1
        char_count_since_last_page = 0
        
        # 1. Scan to find start page (and establish current page state)
        for i in range(start_idx):
             if i >= len(doc.paragraphs): break
             p = doc.paragraphs[i]
             
             old_page = current_page
             current_page = self._update_page_num(p, current_page, char_count_since_last_page)
             
             if current_page > old_page:
                 char_count_since_last_page = 0
             else:
                 # Heuristic Fallback scan
                 p_len = len(p.text)
                 if p_len > 0:
                     char_count_since_last_page += p_len
                 else:
                     char_count_since_last_page += 50 # Empty line weight
                     
                 if char_count_since_last_page > 800: # Consistent with extraction loop
                     current_page += 1
                     char_count_since_last_page = 0
             
        content.append(f"[第 {current_page} 页]")
        
        # 2. Extract Range
        char_count_since_last_page = 0 
        total_extracted_chars = 0
        
        for i in range(start_idx, end_idx + 1):
            if i >= len(doc.paragraphs): break
            
            # Stop if limit reached
            if total_extracted_chars > limit:
                content.append("\n[...该章节内容过长已只截取部分...]")
                break
                
            p = doc.paragraphs[i]
            
            # Check for page num update
            old_page = current_page
            current_page = self._update_page_num(p, current_page, char_count_since_last_page)
            
            page_changed = False
            if current_page > old_page:
                page_changed = True
                char_count_since_last_page = 0
                content.append(f"\n\n[第 {current_page} 页]")
            
            text = p.text.strip()
            if text:
                # Heuristic check
                if not page_changed:
                    char_count_since_last_page += len(text)
                    if char_count_since_last_page > 800: # Tuned down from 1200 for Chinese density
                         current_page += 1
                         content.append(f"\n\n[第 {current_page} 页 (估算)]")
                         char_count_since_last_page = 0
                
                content.append(text)
                total_extracted_chars += len(text)
            else:
                # Empty paragraph = Vertical spacing
                # Treat as approx 50 chars of vertical space
                if not page_changed:
                    char_count_since_last_page += 50
                    if char_count_since_last_page > 800:
                        current_page += 1
                        content.append(f"\n\n[第 {current_page} 页 (估算)]")
                        char_count_since_last_page = 0
                content.append("\n") # Preserve visual spacing
                
        return "\n".join(content)

    def _roman_to_int(self, s):
        s = s.upper()
        romans = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100}
        total = 0
        prev_value = 0
        for char in reversed(s):
            value = romans.get(char, 0)
            if value < prev_value:
                total -= value
            else:
                total += value
            prev_value = value
        return total

    def _update_page_num(self, para, current_page, char_count_since_last_page=0):
        # A. Paragraph Property Check (Page Break Before)
        if para.paragraph_format.page_break_before:
             logging.info(f"[Page Debug] Page Break Before detected. {current_page} -> {current_page + 1}")
             return current_page + 1

        # B. XML Check (Highest Priority - Physical Breaks)
        # 1. Standard Breaks
        for run in para.runs:
            run_xml = run._element.xml
            if 'w:lastRenderedPageBreak' in run_xml or ('w:br' in run_xml and 'w:type="page"' in run_xml):
                # logging.info(f"[Page Debug] XML Physical Break detected. {current_page} -> {current_page + 1}")
                return current_page + 1
        
        # 2. Field Codes (Detect 'PAGE' field cached result)
        # Inspired by User Suggestion: Detect <w:instrText>PAGE</w:instrText>
        # and looking for the cached result in <w:t>
        try:
            if 'PAGE' in para._element.xml:
                # Naive XML string check first to avoid expensive parsing if not present
                xml_str = para._element.xml
                
                # CRITICAL FIX: "NUMPAGES" contains "PAGE", causing False Positives (e.g. Page 1 of 283 -> Jumps to 283)
                if 'NUMPAGES' in xml_str or 'SECTIONPAGES' in xml_str:
                     pass # Ignore Total Page counts
                elif '<w:instrText>PAGE</w:instrText>' in xml_str or 'PAGE' in xml_str:
                     # Attempt to find the number in the XML text (this often holds the cached value)
                     # Exclude "NUMPAGES" matches explicitly in regex
                     # Match "PAGE" whole word logic is hard in XML string, so reliance on the 'if' above is key.
                     # But verify regex doesn't catch remnants.
                     
                     # Simple check: Ensure we are matching the PAGE instruction result
                     # Regex: look for PAGE instr, then find result.
                     # This is heuristic.
                     field_res_match = re.search(r'PAGE\s*<.*?>(\d+)</w:t>', xml_str, re.DOTALL) 
                     if field_res_match:
                         cached_num = int(field_res_match.group(1))
                         # Add sanity check: Page jump shouldn't be TOO massive from current?
                         # e.g. if current=5, detected=283. Likely wrong.
                         # Unless it's a huge doc. 
                         # But let's trust the NUMPAGES filter first.
                         
                         if 0 < cached_num < 5000 and cached_num > current_page:
                             # Safe Jump Guard: preventing massive jumps (e.g. > 50 pages) unless verified?
                             # For now, just logging.
                             logging.info(f"[Page Debug] Field Code 'PAGE' result found: {cached_num}")
                             return cached_num
        except:
            pass
        
        p_text = para.text.strip()
        if not p_text: return current_page

        # B. Visual Marker
        if "分页符" in p_text and "....." in p_text:
             logging.info(f"[Page Debug] Visual Marker detected. {current_page} -> {current_page + 1}")
             return current_page + 1
        
        # C. Footer (Enhanced with Roman & Continuous Logic)
        detected_num = None
        
        # 1. Roman Numerals
        # Flexible match: "第·II·页", "- II -", "II", "第 II 页"
        # Match Roman between non-word chars or specific separators
        roman_match = re.search(r'^\s*(?:第|[-—·\.\u2022\u2027\u30fb])*\s*([IVXivx]+)\s*(?:页|[-—·\.\u2022\u2027\u30fb])*\s*[.。]?$', p_text)
        if roman_match:
             try:
                 r_str = roman_match.group(1)
                 if len(r_str) < 10: 
                     detected_num = self._roman_to_int(r_str)
                     # logging.info(f"[Page Debug] Roman Numeral Footer found: '{p_text}' -> {detected_num}")
             except:
                 pass

        # 2. Standard Digits
        if detected_num is None:
            patterns = [
                r'^\s*(?:page|第|[·\.\-_])?\s*(\d+)\s*(?:page|页|[·\.\-_])?\s*[.。]?$', 
                # Broaden "Decorative" support: ANY non-digit separator between 第 and Number, Number and 页
                r'^\s*第\D*(\d+)\D*页\s*[.。]?$',
                r'^\s*[-—]\s*(\d+)\s*[-—]\s*$', 
                r'^\s*(\d+)\s*/\s*\d+\s*$' 
            ]
            for pat in patterns:
                pm = re.search(pat, p_text, re.IGNORECASE)
                if pm:
                    try:
                        detected_num = int(pm.group(1))
                        # logging.info(f"[Page Debug] Digit Footer found: '{p_text}' -> {detected_num}")
                        break 
                    except:
                        pass
        
        # D. Continuous Logic Application
        if detected_num is not None:
            if 0 < detected_num < 5000:
                # 1. Sync Forward
                if detected_num > current_page:
                    logging.info(f"[Page Debug] Sync Forward: Footer '{p_text}' says {detected_num}, current is {current_page}. Jumping to {detected_num}.")
                    return detected_num
                
                # 2. Restart/Implicit Break
                if char_count_since_last_page > 50:
                    # If detected <= current, it means a restart. 
                    # If we have lots of chars, we missed a break.
                    logging.info(f"[Page Debug] Implicit Break: Footer '{p_text}' says {detected_num} (<= {current_page}), but char buffer {char_count_since_last_page} > 50. Incrementing to {current_page + 1}.")
                    return current_page + 1

        return current_page

    def get_all_content(self, limit=None) -> str:
        """
        Retrieves the full content of the current main document with page numbers.
        """
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc: return ""
        
        # Reuse extract logic but for the whole range
        total_paras = len(target_doc.paragraphs)
        text_content = self._extract_range_with_pages(target_doc, 0, total_paras, limit=limit or 100000)
        
        # Append Table Content
        table_content = []
        if target_doc.tables:
            table_content.append("\n\n[表格数据 (Table Data)]")
            for i, table in enumerate(target_doc.tables):
                table_content.append(f"\n[Table {i+1}]")
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    # Simple CSV-like format
                    table_content.append(" | ".join(row_cells))
        
        return text_content + "\n".join(table_content)

    def get_relevant_reference_context(self, user_query: str) -> str:
        """
        Smart Context Retrieval (Universal Adaptive Strategy).
        1. Short Docs (< 300 paras): Return FULL content.
        2. Long Docs: Coarse-to-Fine Search.
        """
        if not self.reference_docs:
            return ""
            
        relevant_parts = []
        import re
        # Split by space, comma, newline for rudimentary keyword matching
        raw_keywords = re.split(r'[\s,，。]+', user_query.strip())
        keywords = [k for k in raw_keywords if len(k) > 1]
        
        matches_found = False
        
        for ref in self.reference_docs:
            doc = ref['doc']
            filename = ref['filename']
            para_count = len(doc.paragraphs)
            
            # --- ADAPTIVE STRATEGY: Small Doc Optimization ---
            if para_count < 300:
                # Direct Injection for small docs
                full_text = self._extract_range_with_pages(doc, 0, para_count, limit=50000)
                relevant_parts.append(f"--- 参考文档：{filename} (全文) ---\n{full_text}\n")
                matches_found = True
                continue
            # -------------------------------------------------
            
            # 1. Build Outline
            outline = []
            for i, para in enumerate(doc.paragraphs):
                if not para.style: continue
                s_name = para.style.name.lower()
                if "heading" in s_name or "title" in s_name:
                     outline.append({"index": i, "text": para.text.strip()})
            outline.append({"index": len(doc.paragraphs), "text": "END"})
            
            # 2. Identify Sections to Extract
            sections_to_extract = []
            for i in range(len(outline) - 1):
                heading = outline[i]
                next_heading = outline[i+1]
                
                is_hit = False
                for kw in keywords:
                    if kw.lower() in heading["text"].lower():
                        is_hit = True
                        break
                
                if is_hit:
                    sections_to_extract.append((heading["index"], next_heading["index"], heading["text"]))
            
            # UPGRADE: Deep Semantic Search (Full Body Scan)
            # If no Heading matches found, scan the entire text for the keywords.
            if not sections_to_extract:
                hits = []
                for kw in keywords:
                    # Scan all paragraphs
                    for i, para in enumerate(doc.paragraphs):
                        if kw.lower() in para.text.lower():
                            hits.append(i)
                            if len(hits) > 5: break # Cap matches for performance
                    if hits: break
                
                if hits:
                    # Consolidate hits into windows
                    # Window: -20 to +50 paragraphs around the hit
                    matches_found = True
                    for hit_idx in hits:
                        start_window = max(0, hit_idx - 20)
                        end_window = min(len(doc.paragraphs), hit_idx + 80)
                        sections_to_extract.append((start_window, end_window, f"内容深度匹配 (段落 {hit_idx})"))
            
            if sections_to_extract:
                matches_found = True
                relevant_parts.append(f"--- 参考文档：{filename} (精选章节) ---")
                
                # 3. Fast Page Scan (Build Page Map)
                # Map paragraph index -> Page Number at the start of that paragraph
                page_map = {}
                current_page = 1
                for i, para in enumerate(doc.paragraphs):
                    page_map[i] = current_page
                    
                    # A. XML Check
                    for run in para.runs:
                        run_xml = run._element.xml
                        if 'w:lastRenderedPageBreak' in run_xml or ('w:br' in run_xml and 'w:type="page"' in run_xml):
                            current_page += 1
                    
                    p_text = para.text.strip()
                    if not p_text: continue

                    # B. Visual Marker (Explicit Page Break text)
                    if "分页符" in p_text and "....." in p_text:
                         current_page += 1
                    
                    # C. Page Number Sync (Footer Detection)
                    # Support formats: "57", "第 3 页", "第-3-页", "第·3·页", "Page 57", "- 57 -"
                    pm = re.search(r'^\s*(?:page|第|[·\.\-_])?\s*(\d+)\s*(?:page|页|[·\.\-_])?\s*$', p_text, re.IGNORECASE)
                    if pm:
                        try:
                            detected_num = int(pm.group(1))
                            if 0 < detected_num < 2000 and len(p_text) < 20:
                                current_page = detected_num
                        except:
                            pass
                
                # 4. Extract Content
                # Strategy: Distribute context window. Cap each section at ~2000 chars.
                # If we have many matches, this ensures diversity.
                MAX_SECTION_CHARS = 2000
                
                for start, end, title in sections_to_extract:
                    section_content = ""
                    char_count = 0
                    
                    # Get start page from map
                    start_page = page_map.get(start, 1)
                    current_section_page = start_page
                    section_content += f"[第 {start_page} 页] "
                    
                    for p_idx in range(start, end):
                        if char_count > MAX_SECTION_CHARS:
                            section_content += "\n[...该章节内容过长已只截取部分...]"
                            break
                        
                        para = doc.paragraphs[p_idx]
                        txt = para.text.strip()
                        
                        # Check for page break WITHIN the section to update marker
                        # Use the pre-calculated map? 
                        # If page_map[p_idx] > current_section_page, insert marker
                        this_para_page = page_map.get(p_idx, 1)
                        if this_para_page > current_section_page:
                            current_section_page = this_para_page
                            section_content += f"\n\n[第 {current_section_page} 页]\n"
                        
                        if txt:
                            section_content += txt + "\n"
                            char_count += len(txt)
                            
                    relevant_parts.append(f"\n### 章节：{title} ###\n{section_content}\n")
            
            else:
                # No matches in this file
                pass

        if not matches_found:
             # Fallback
             return self.get_reference_context()
             
        return "\n".join(relevant_parts)

    def extract_with_images(self, docx_path, image_output_dir):
        """
        Extract text and images from DOCX.
        Images are saved to image_output_dir.
        Returns markdown text with image links pointing to relative paths.
        """
        if not os.path.exists(docx_path):
             return "File not found."
             
        if not os.path.exists(image_output_dir):
            os.makedirs(image_output_dir, exist_ok=True)

        return self._convert_to_markdown_with_images(docx_path, image_output_dir)

    def _convert_to_markdown_with_images(self, docx_path, image_output_dir):
        import mammoth
        
        def convert_image(image):
            with image.open() as image_bytes:
                content = image_bytes.read()
                # Generate unique filename
                import uuid
                filename = f"uploaded_image_{uuid.uuid4()}.png"
                filepath = os.path.join(image_output_dir, filename)
                
                with open(filepath, "wb") as f:
                    f.write(content)
                
                # Return dictionary with src (URL path)
                # Frontend should proxy /static/images/ to backend static folder
                return {"src": f"/static/images/{filename}", "alt": "Uploaded Image"}

        try:
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(
                    docx_file, 
                    convert_image=convert_image  # Pass directly, my handler returns {src: ...}
                )
                return result.value
        except Exception as e:
            return f"Error extracting with images: {str(e)}"

    def get_preview_data(self, start: int = 0, limit: int = None) -> List[Dict[str, Any]]:
        """
        Extracts paragraphs to send to frontend for preview.
        Returns a list of paragraphs, each containing a list of runs with formatting.
        Supports pagination via start and limit.
        """
        # Use staging doc if it exists, otherwise the committed doc
        target_doc = self.staging_doc if self.staging_doc else self.doc
        
        if not target_doc:
            return []
        
        preview_data = []
        
        # Paginate
        paragraphs = target_doc.paragraphs
        if limit:
            paragraphs = paragraphs[start : start + limit]
        else:
            paragraphs = paragraphs[start:]
            
        for i, para in enumerate(paragraphs):
            # Calculate absolute index
            abs_index = start + i
            runs_data = []
            for run in para.runs:
                # Extract color as hex string if exists
                color = None
                if run.font.color and run.font.color.rgb:
                    color = f"#{run.font.color.rgb}"
                
                runs_data.append({
                    "text": run.text,
                    "bold": bool(run.bold),
                    "italic": bool(run.italic),
                    "underline": bool(run.underline),
                    "color": color,
                    "fontSize": run.font.size.pt if run.font.size else None
                })

            preview_data.append({
                "id": i,
                "text": para.text, # Keep raw text for easy reading/debugging
                "style": para.style.name if para.style else "Normal",
                "runs": runs_data
            })
        return preview_data
    def get_html_preview(self, start: int = 0, limit: int = None) -> str:
        """
        Generates an HTML representation of the document with injected IDs.
        Supports pagination via start and limit.
        """
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc:
            return "<div class='empty-doc'>No document loaded</div>"

        html_parts = ['<div class="docx-viewer">']
        
        # We need a unified index for "paragraphs" to match the AI's understanding
        # IF we change to block iteration, indices might shift if we count tables as blocks.
        # For now, simplistic paragraph iteration for large doc optimization.
        
        paragraphs = target_doc.paragraphs
        total_paras = len(paragraphs)
        
        if limit:
            slice_end = min(total_paras, start + limit)
            # This slicing is tricky if we mix types. 
            # We need to collect all blocks first? That might be slow.
            # Generator slicing?
            # Let's iterate and count.
            pass
        else:
            slice_end = total_paras

        current_index = 0
        rendered_count = 0
        
        for block in self.iter_block_items(target_doc):
            if current_index < start:
                current_index += 1
                continue
            
            if limit and rendered_count >= limit:
                break

            abs_index = current_index
            
            if isinstance(block, Paragraph):
                html_parts.append(self._render_paragraph_html(block, abs_index))
            elif isinstance(block, Table):
                html_parts.append(self._render_table_html(block, abs_index))
            
            current_index += 1
            rendered_count += 1

        
        if slice_end < total_paras:
             html_parts.append(f"<div class='page-break'>... ({total_paras - slice_end} more paragraphs) ...</div>")
             
        html_parts.append('</div>')
        html_parts.append('</div>')
        return "".join(html_parts)

    def get_document_structure(self) -> list:
        """
        Scans the document for Headings and returns a TOC structure.
        Uses Outline levels for robustness.
        """
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc:
            return []
        return self._extract_toc_from_doc(target_doc)

    def _extract_toc_from_doc(self, doc) -> list:
        toc = []
        # Re-use our robust page counting logic!
        # Don't use naive (i // 100).
        current_page = 1
        char_count_since_last_page = 0
        
        for i, para in enumerate(doc.paragraphs):
            # Update Page Count (Sync with _extract_range_with_pages)
            new_page = self._update_page_num(para, current_page, char_count_since_last_page)
            if new_page > current_page:
                current_page = new_page
                char_count_since_last_page = 0
            
            # Update char count heuristic
            p_text = para.text.strip()
            if p_text:
                char_count_since_last_page += len(p_text)
                if char_count_since_last_page > 800: # Consistent with main loop
                     current_page += 1
                     char_count_since_last_page = 0
            else:
                # Consistent with main loop empty line heuristic
                char_count_since_last_page += 50
                if char_count_since_last_page > 800:
                     current_page += 1
                     char_count_since_last_page = 0

            level = 0
            text = p_text # Use the stripped text we already got
            if not text: continue

            # 1. Check Outline Level (Robust)
            try:
                # outline_level is an integer 0-8 for Level 1-9. 9 means Body Text.
                # But python-docx might raise error if not present? No, usually valid.
                lvl = para.paragraph_format.outline_level
                if 0 <= lvl <= 8:
                    level = lvl + 1
            except:
                pass
            
            # 2. Fallback to Style Name if outline level is Body Text (9)
            if level == 0 and para.style:
                style_name = para.style.name.lower()
                if "heading 1" in style_name: level = 1
                elif "heading 2" in style_name: level = 2
                elif "heading 3" in style_name: level = 3
                elif "heading 4" in style_name: level = 4
                elif "title" in style_name: level = 1
            
            # 3. Fallback to Regex (Heuristic for poorly formatted docs)
            # 3. Fallback to Regex (Heuristic for poorly formatted docs)
            if level == 0:
                # Optimized Heuristic Patterns:
                # Level 1: "一、", "第一章", "1、" (if bold), "1. " (if bold), "Chapter 1"
                # Level 2: "（一）", "(一)", "1.1", "1.1.1", "1.11"
                
                # Check for bold first (strong signal for unstructured docs)
                is_bold = False
                if para.runs:
                    # Check if ANY run is bold? Or if the majority is bold?
                    # Simple check: if first run is bold or style implies bold
                    for r in para.runs:
                        if r.bold and len(r.text.strip()) > 0:
                            is_bold = True
                            break
                            
                # Matches "一、", "二、" ...
                if re.match(r"^[一二三四五六七八九十]+[、.,，]", text):
                    level = 1
                # Matches "第一章", "第二节"...
                elif re.match(r"^第[一二三四五六七八九十0-9]+[章节]", text):
                    level = 1
                    
                # Matches "（一）", "(一)" - supports Chinese and English brackets
                elif re.match(r"^[（(][一二三四五六七八九十]+[）)]", text):
                    level = 2
                    
                # Matches "1.", "1、" (Single digit headers)
                # Note: Logic priority matters. 
                elif re.match(r"^\d+[、.．]\s?$", text) or  (re.match(r"^\d+[、.．]\s+", text) and len(text) < 30):
                     # "1." or "1、" usually Level 3 if inside (一), but sometimes Level 1 if no Chinese numbers used.
                     # We'll assign Level 3 to distinguish from (一).
                     level = 3
                     
                # Matches "1.1", "1.1.1" (Points/Sub-points)
                elif re.match(r"^\d+(\.\d+)+\s*", text) and len(text) < 50:
                    # Calculate depth based on dots
                    # 1.1 -> 1 dot -> level 4
                    # 1.1.1 -> 2 dots -> level 5
                    dots = text.split(" ")[0].count(".")
                    level = 3 + dots # 1.1 -> 3+1=4, 1.1.1 -> 5
                    if level > 6: level = 6

                # Matches "1 " (Digit space text) - e.g. "1 对项目的理解"
                # This is common in the user's screenshot
                elif re.match(r"^\d+\s+[\u4e00-\u9fa5]", text) and len(text) < 40:
                     level = 3
                     
                # Matches simple bold lines that look like titles
                elif is_bold and len(text) < 40 and not text[-1] in ["。", ".", "”", "！", "!"]:
                     if level == 0: level = 3 # Default to level 3 for unspecified bold headers
            
            if level > 0:
                # snippet Extraction (Smart Outline)
                snippet = ""
                # Peek ahead for snippet
                # We want the first non-empty text that is NOT a title/heading
                start_peek_id = i + 1
                search_limit = 5 # Look at next 5 paragraphs max
                found_snippet = False
                
                for peek_i in range(start_peek_id, min(start_peek_id + search_limit, len(doc.paragraphs))):
                    peek_p = doc.paragraphs[peek_i]
                    peek_text = peek_p.text.strip()
                    if not peek_text:
                        continue
                        
                    # Check if this next para is also a header?
                    # Simple heuristic: if it's very short and bold, or matches our regex, might be subheader
                    # But for snippet, we'll take it unless it's obviously a strong header
                    # If snippet is too short, we might append more? For now, just take the first chunk.
                    snippet = peek_text[:200]
                    if len(snippet) < 200 and len(peek_text) > 200:
                         snippet += "..."
                    break

                toc.append({
                    "id": i,
                    "title": text,
                    "level": level,
                    "page": current_page, # Use the synced robust page number
                    "snippet": snippet # Smart Outline Feature
                })
        
        # Calculate End Indices
        total_paras = len(doc.paragraphs)
        for idx, item in enumerate(toc):
            if idx < len(toc) - 1:
                item["end_id"] = toc[idx + 1]["id"] - 1
            else:
                item["end_id"] = total_paras - 1
                
        return toc

    def get_global_context(self) -> str:
        """
        Generates a compressed summary of the document (Heading + First Paragraph).
        Used to provide global context to the LLM.
        """
        toc = self.get_document_structure()
        if not toc:
            return ""
            
        summary_lines = ["【全书大纲与摘要】:"]
        target_doc = self.staging_doc if self.staging_doc else self.doc
        
        for item in toc:
            # Add Title
            indent = "  " * (item['level'] - 1)
            summary_lines.append(f"{indent}- {item['title']}")
            
            # Add snippet of content (next paragraph if it exists and isn't a heading)
            # This is a heuristic: "Read the first bit of this section"
            start_id = item['id']
            if start_id + 1 < len(target_doc.paragraphs):
                next_para = target_doc.paragraphs[start_id + 1]
                # Avoid if next para is also a heading
                if next_para.text.strip() and (not next_para.style or "heading" not in next_para.style.name.lower()):
                    # Truncate to ~100 chars
                    content_snippet = next_para.text.strip()[:100] + "..."
                    summary_lines.append(f"{indent}  (内容概览: {content_snippet})")
                    
        return "\n".join(summary_lines)

    def _get_images_from_run(self, run, doc):
        """
        Extracts images from a run.
        Returns a list of URLs (local static paths).
        """
        image_urls = []
        try:
            # Namespaces
            nsmap = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'p': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            }
            
            # 1. Search for blip (Bitmap)
            # This covers typical inserted images
            blips = run.element.findall('.//a:blip', nsmap)
            for blip in blips:
                rId = blip.get(f"{{{nsmap['r']}}}embed") # Get r:embed attribute
                if rId:
                    image_part = doc.part.related_parts.get(rId)
                    if image_part:
                        url = self._save_image_part(image_part)
                        if url: image_urls.append(url)
            
            # 2. Search for v:imagedata (Legacy VML) - if lxml used different map
            # Fallback if needed, but 'a:blip' covers most docx.
            
        except Exception as e:
            # logging.error(f"Image extraction error: {e}")
            pass
            
        return image_urls

    def _save_image_part(self, image_part):
        try:
            import uuid
            # Determine content type / extension
            content_type = image_part.content_type
            ext = "png"
            if "jpeg" in content_type: ext = "jpg"
            elif "bmp" in content_type: ext = "bmp"
            elif "gif" in content_type: ext = "gif"
            
            filename = f"extracted_{uuid.uuid4()}.{ext}"
            
            # Ensure dir exists
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            static_img_dir = os.path.join(base_dir, 'static', 'images')
            if not os.path.exists(static_img_dir):
                os.makedirs(static_img_dir, exist_ok=True)
                
            filepath = os.path.join(static_img_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(image_part.blob)
                
            return f"/static/images/{filename}"
        except Exception as e:
             return None
        
    def get_paragraph_count(self):
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if target_doc:
            return len(target_doc.paragraphs)
        return 0


    def iter_block_items(self, parent):
        """
        Yields each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph.
        """
        # Check if parent has access to the body element (duck typing)
        if hasattr(parent, 'element') and hasattr(parent.element, 'body'):
            parent_elm = parent.element.body
        else:
             return 

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _render_paragraph_html(self, para, abs_index):
        text = para.text
        style_class = "p-normal"
        if para.style:
            s_name = para.style.name.lower()
            if "heading 1" in s_name: style_class = "p-h1"
            elif "heading 2" in s_name: style_class = "p-h2"
            elif "heading 3" in s_name: style_class = "p-h3"
            elif "title" in s_name: style_class = "p-title"
            elif "list" in s_name: style_class = "p-list"

        inner_html = ""
        if not para.runs:
            inner_html = text 
        else:
            for run in para.runs:
                # Images
                image_urls = self._get_images_from_run(run, para.part) # Pass part? Doc?
                # accessing doc from run is not direct in all versions. 
                # run.part might be enough. _get_images_from_run uses run.element
                # Actually _get_images_from_run signature used doc.
                # Let's verify _get_images_from_run usage.
                # It accesses doc.part.related_parts.
                # Here we might need the main doc reference if not attached to run.
                
                # Correction: run.part is robust.
                # But my _get_images_from_run implementation used `doc` argument.
                # I should update _get_images_from_run to accept `part` or use `self.doc`.
                # Let's pass `self.staging_doc` or `self.doc` to helper if needed.
                target_doc = self.staging_doc if self.staging_doc else self.doc
                image_urls = self._get_images_from_run(run, target_doc)
                
                for img_url in image_urls:
                    inner_html += f'<img src="{img_url}" style="max-width:100%; height:auto; display:block; margin: 10px auto;" />'

                run_text = run.text
                if not run_text and not image_urls: continue
                    
                run_text = run_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                
                span_style = ""
                if run.bold: span_style += "font-weight:bold;"
                if run.italic: span_style += "font-style:italic;"
                if run.underline: span_style += "text-decoration:underline;"
                if run.font.color and run.font.color.rgb:
                     span_style += f"color:#{run.font.color.rgb};"
                
                if span_style:
                    inner_html += f"<span style='{span_style}'>{run_text}</span>"
                else:
                    inner_html += run_text

        if not inner_html: inner_html = "&nbsp;"
        return f'<div class="doc-para {style_class}" data-id="{abs_index}">{inner_html}</div>'

    def _render_table_html(self, table, abs_index):
        rows_html = []
        for row in table.rows:
            cells_html = []
            for cell in row.cells:
                # Recursively render cell content? Or just plain text?
                # For simplicity, join paragraphs
                cell_text = []
                for p in cell.paragraphs:
                    cell_text.append(p.text)
                content = "<br>".join(cell_text) if cell_text else "&nbsp;"
                cells_html.append(f"<td style='border:1px solid #ccc; padding:4px;'>{content}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
            
        return f'<div class="doc-table" data-id="{abs_index}"><table style="border-collapse:collapse; width:100%; border:1px solid #777;">{"".join(rows_html)}</table></div>'

    def create_staging_copy(self):
        """
        Creates a deep copy of the current document for staging.
        We do this by saving to a memory stream and reloading.
        """
        if not self.doc:
            return
        
        stream = io.BytesIO()
        self.doc.save(stream)
        stream.seek(0)
        self.staging_doc = Document(stream)

    def commit_staging(self):
        """
        Makes the staging document the permanent one.
        """
        if self.staging_doc:
            self.doc = self.staging_doc
            self.staging_doc = None
            return True
        return False

    def discard_staging(self):
        """
        Discards the staging document.
        """
        self.staging_doc = None
        return True

    def apply_patch(self, patch_json: List[Dict[str, Any]], use_staging: bool = False) -> bool:
        """
        Applies a list of surgical edits.
        """
        target_doc = self.staging_doc if use_staging else self.doc
        
        if not target_doc:
            return False

        for op in patch_json:
            op_type = op.get("op")
            if op_type == "replace_text":
                self._replace_paragraph_text(target_doc, op.get("target_id"), op.get("new_text"))
            elif op_type == "search_replace":
                self._search_replace(target_doc, op.get("target_id"), op.get("find_text"), op.get("replace_text"))
            # Future: Add insert_paragraph, delete_paragraph, etc.
        
        return True

    def execute_code(self, code: str) -> bool:
        """
        Executes the provided Python code on the document.
        The code has access to 'doc' (the Document object) and 'docx' (the module).
        """
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc:
            return False, "No document loaded"

        # Define the execution context
        # Common imports for docx manipulation
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        import re
        
        
        # Mongo-patch Paragraph to support insert_paragraph_after
        # This fixes the AI hallucination issue
        def insert_paragraph_after(self, text=None, style=None):
            """Insert a new paragraph after this one."""
            new_p = self._parent.add_paragraph(text, style)
            self._element.addnext(new_p._element)
            return new_p
            
        if not hasattr(Paragraph, 'insert_paragraph_after'):
            Paragraph.insert_paragraph_after = insert_paragraph_after

        # Monkey-patch Paragraph to have an 'id' property (returning index)
        # This fixes AI assuming paragraphs have IDs matching the context
        if not hasattr(Paragraph, 'id'):
            @property
            def paragraph_id(self):
                try:
                    # Attempt to find index in parent document
                    # This is slow O(N) but functional
                    return self._parent.paragraphs.index(self)
                except ValueError:
                    return -1
            Paragraph.id = paragraph_id

        # Helper to insert image
        def insert_image(paragraph, image_path_or_url, width=None):
            """
            Inserts an image after the given paragraph.
            Handles /static/images/ paths by resolving to local file system.
            Enforces SINGLE line spacing for the image paragraph.
            """
            try:
                # Resolve path
                img_path = image_path_or_url
                if '/static/images/' in img_path:
                   filename = img_path.split('/static/images/')[-1]
                   import os
                   base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                   img_path = os.path.join(base_dir, 'backend', 'static', 'images', filename)
                
                if not os.path.exists(img_path):
                    paragraph.insert_paragraph_after(f"[Image not found: {img_path}]")
                    return

                # Create run
                run = paragraph.add_run()
                if width:
                    run.add_picture(img_path, width=width)
                else:
                    run.add_picture(img_path, width=Cm(15)) # Default width
                
                # Format paragraph: Center + Single Line Spacing
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                return True
            except Exception as e:
                paragraph.insert_paragraph_after(f"[Error inserting image: {str(e)}]")
                return False

        # Monkey-patch Run.add_picture to handle /static/images/ paths automatically
        # This is the "Safety Net" for when AI uses standard .add_picture() instead of insert_image()
        from docx.text.run import Run
        original_add_picture = Run.add_picture

        def patched_add_picture(self, image_path_or_stream, width=None, height=None):
            if isinstance(image_path_or_stream, str) and '/static/images/' in image_path_or_stream:
                try:
                    filename = image_path_or_stream.split('/static/images/')[-1]
                    import os
                    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                    real_path = os.path.join(base_dir, 'backend', 'static', 'images', filename)
                    if os.path.exists(real_path):
                        image_path_or_stream = real_path
                except Exception as e:
                    logging.error(f"Failed to resolve image path in patched_add_picture: {e}")
            
            return original_add_picture(self, image_path_or_stream, width, height)

        Run.add_picture = patched_add_picture

        # MONKEY PATCH 2: Enable insert_paragraph_after for Agents
        def insert_paragraph_after(self, text=None, style=None):
            """
            Inserts a paragraph after this paragraph.
            """
            # Find the next sibling that is a paragraph, or just insert after in XML
            from docx.oxml.text.paragraph import CT_P
            
            new_p = self._parent.add_paragraph(text, style)
            # Move it to after self
            self._parent._element.remove(new_p._element)
            self._element.addnext(new_p._element)
            return new_p

        Paragraph.insert_paragraph_after = insert_paragraph_after

        local_scope = {
            "doc": target_doc,
            "Document": Document,
            "Paragraph": Paragraph,
            "RGBColor": RGBColor,
            "Pt": Pt,
            "Inches": Inches,
            "Cm": Cm,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_LINE_SPACING": WD_LINE_SPACING, # Expose constant
            "qn": qn,
            "re": re,
            "print": print,
            "smart_replace": self.smart_replace, 
            "flexible_replace": self.flexible_replace, 
            "search_replace": self.search_replace, 
            "apply_markdown": self.replace_with_markdown, 
            "set_east_asian_font": self.set_east_asian_font,
            "insert_image": insert_image 
        }
        
        # FIX: The LLM sometimes hallucinates 'import utils' or 'utils.smart_replace'.
        # We handle this by stripping the import and providing a mock object.
        if "import utils" in code or "from utils" in code:
             # Remove the import line to avoid ModuleNotFoundError
             code = re.sub(r'^\s*import utils.*$', '', code, flags=re.MULTILINE)
             code = re.sub(r'^\s*from utils import.*$', '', code, flags=re.MULTILINE)
             
             class MockUtils:
                 pass
             mock_utils = MockUtils()
             # Map common methods to the mock object
             mock_utils.smart_replace = self.smart_replace
             mock_utils.flexible_replace = self.flexible_replace
             mock_utils.search_replace = self.search_replace
             mock_utils.apply_markdown = self.replace_with_markdown
             mock_utils.insert_image = insert_image
             mock_utils.set_east_asian_font = self.set_east_asian_font
             
             local_scope["utils"] = mock_utils

        try:
            exec(code, local_scope)
            return True, ""
        except Exception as e:
            error_msg = f"Error executing code: {str(e)}"
            logging.error(error_msg)
            return False, error_msg

    def flexible_replace(self, doc, find_text: str, replace_text: str) -> bool:
        """
        Tries smart_replace first (preserving formatting).
        If that fails to find/replace anything, falls back to raw text replacement (prioritizing content).
        """
        if not find_text:
            logging.warning("flexible_replace called with empty find_text")
            return False

        logging.info(f"flexible_replace: find='{find_text[:50]}...', replace='{replace_text[:50]}...'")

        # 1. Try smart replace
        if self.smart_replace(doc, find_text, replace_text):
            logging.info("flexible_replace: smart_replace succeeded")
            return True
        
        logging.info("flexible_replace: smart_replace failed, trying raw replace")
        
        # 2. Fallback: Raw text replacement
        # This might lose formatting for the specific paragraph, but ensures content is updated.
        replaced = False
        for para in doc.paragraphs:
            if find_text in para.text:
                # Direct assignment to para.text clears runs and sets new text
                try:
                    para.text = para.text.replace(find_text, replace_text)
                    replaced = True
                    logging.info(f"flexible_replace: raw replace succeeded in paragraph {i}")
                except Exception as e:
                    logging.error(f"flexible_replace: raw replace failed: {e}")
        
        if not replaced:
            logging.warning("flexible_replace: text not found in any paragraph")
            
        return replaced

    def smart_replace(self, doc, find_text: str, replace_text: str) -> bool:
        """
        Replaces text while trying to preserve formatting.
        Handles cases where text is split across runs.
        Strategy: Find the span of runs, replace text in the first run, clear others.
        """
        replaced = False
        for para in doc.paragraphs:
            if find_text in para.text:
                # Naive check: if it's in one run, just do it
                replaced = False
                for run in para.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        replaced = True
                        break
                
                if replaced:
                    continue
                
                # Complex check: Text is split across runs
                # This is a hard problem to solve perfectly without a proper interval tree or character map.
                # Simplified approach:
                # 1. Build a map of character indices to run indices.
                # 2. Find start/end char index of the match.
                # 3. Identify start/end runs.
                
                text = para.text
                start_index = text.find(find_text)
                if start_index == -1:
                    continue
                    
                end_index = start_index + len(find_text)
                
                # Map chars to runs
                current_pos = 0
                start_run_idx = -1
                end_run_idx = -1
                
                for i, run in enumerate(para.runs):
                    run_len = len(run.text)
                    run_end = current_pos + run_len
                    
                    if start_run_idx == -1 and start_index < run_end:
                        start_run_idx = i
                        # Calculate offset in this run
                        start_offset = start_index - current_pos
                    
                    if end_run_idx == -1 and end_index <= run_end:
                        end_run_idx = i
                        end_offset = end_index - current_pos
                        break
                    
                    current_pos += run_len
                
                if start_run_idx != -1 and end_run_idx != -1:
                    # We found the span
                    # Run[start_run_idx] ... Run[end_run_idx]
                    
                    # 1. Update Start Run
                    # We want to keep: prefix + replacement + suffix (if end run is same)
                    # But if end run is different, we keep prefix + replacement.
                    
                    start_run = para.runs[start_run_idx]
                    prefix = start_run.text[:start_index - (current_pos - len(start_run.text) if start_run_idx == end_run_idx else self._get_run_start_pos(para, start_run_idx))]
                    # Wait, calculating offsets again is messy.
                    
                    # Let's use the simpler logic:
                    # Reconstruct the text of the start run to include the replacement.
                    # Clear the text of intermediate runs.
                    # Adjust the text of the end run to keep the suffix.
                    
                    # Re-calculate offsets properly
                    curr = 0
                    for j in range(start_run_idx):
                        curr += len(para.runs[j].text)
                    
                    start_offset_in_run = start_index - curr
                    
                    curr = 0
                    for j in range(end_run_idx):
                        curr += len(para.runs[j].text)
                    end_offset_in_run = end_index - curr
                    
                    if start_run_idx == end_run_idx:
                        # Same run (handled by naive loop usually, but maybe not if multiple occurrences?)
                        run = para.runs[start_run_idx]
                        run.text = run.text[:start_offset_in_run] + replace_text + run.text[end_offset_in_run:]
                    else:
                        # Multi-run
                        # Start run: keep prefix + replace_text
                        para.runs[start_run_idx].text = para.runs[start_run_idx].text[:start_offset_in_run] + replace_text
                        
                        # Middle runs: clear
                        for k in range(start_run_idx + 1, end_run_idx):
                            para.runs[k].text = ""
                            
                        # End run: keep suffix
                        para.runs[end_run_idx].text = para.runs[end_run_idx].text[end_offset_in_run:]
                        replaced = True
        
        return replaced

    def _get_run_start_pos(self, para, run_idx):
        pos = 0
        for i in range(run_idx):
            pos += len(para.runs[i].text)
        return pos

    def _replace_paragraph_text(self, doc, index: int, new_text: str):
        # Deprecated in favor of execute_code, but kept for compatibility if needed
        pass

    def search_replace(self, doc, find_text: str, replace_text: str) -> bool:
        """
        Robust search and replace that handles multiple occurrences and preserves formatting.
        """
        changed_any = False
        for para in doc.paragraphs:
            if find_text not in para.text:
                continue
                
            # Loop until no more matches in this paragraph
            # We need to be careful about infinite loops if replace_text contains find_text
            # But usually we just process the paragraph once or iteratively.
            # Since modifying runs changes indices, it's safer to restart the search in the paragraph after a modification,
            # OR use a while loop.
            
            # Simple safety: limit iterations per paragraph
            iterations = 0
            while find_text in para.text and iterations < 20:
                # Try to find and replace ONE occurrence
                replaced = self._smart_replace_single(para, find_text, replace_text)
                if replaced:
                    changed_any = True
                    iterations += 1
                else:
                    # If text is there but we failed to replace it (e.g. complex split?), break to avoid infinite loop
                    # Or maybe fall back to raw replacement for just this paragraph?
                    # Let's try to be robust.
                    break
        return changed_any

    def _smart_replace_single(self, para, find_text: str, replace_text: str) -> bool:
        """
        Helper to replace the FIRST occurrence of find_text in a paragraph.
        Returns True if replaced.
        """
        # 1. Naive check: if it's in one run
        for run in para.runs:
            if find_text in run.text:
                run.text = run.text.replace(find_text, replace_text, 1)
                return True
        
        # 2. Split across runs
        text = para.text
        start_index = text.find(find_text)
        if start_index == -1:
            return False
            
        end_index = start_index + len(find_text)
        
        # Map chars to runs
        current_pos = 0
        start_run_idx = -1
        end_run_idx = -1
        
        for i, run in enumerate(para.runs):
            run_len = len(run.text)
            run_end = current_pos + run_len
            
            if start_run_idx == -1 and start_index < run_end:
                start_run_idx = i
                # Calculate offset in this run
                # start_offset = start_index - current_pos
            
            if end_run_idx == -1 and end_index <= run_end:
                end_run_idx = i
                # end_offset = end_index - current_pos
                break
            
            current_pos += run_len
        
        if start_run_idx != -1 and end_run_idx != -1:
            # Calculate offsets
            curr = 0
            for j in range(start_run_idx):
                curr += len(para.runs[j].text)
            start_offset_in_run = start_index - curr
            
            curr = 0
            for j in range(end_run_idx):
                curr += len(para.runs[j].text)
            end_offset_in_run = end_index - curr
            
            if start_run_idx == end_run_idx:
                # Same run (handled by naive loop usually)
                run = para.runs[start_run_idx]
                run.text = run.text[:start_offset_in_run] + replace_text + run.text[end_offset_in_run:]
            else:
                # Multi-run
                # Start run: keep prefix + replace_text
                para.runs[start_run_idx].text = para.runs[start_run_idx].text[:start_offset_in_run] + replace_text
                
                # Middle runs: clear
                for k in range(start_run_idx + 1, end_run_idx):
                    para.runs[k].text = ""
                    
                # End run: keep suffix
                para.runs[end_run_idx].text = para.runs[end_run_idx].text[end_offset_in_run:]
            
            return True
            
        return False

    def _search_replace(self, doc, target_id: int, find_text: str, replace_text: str):
        # Deprecated
        pass

    def replace_with_markdown(self, doc, target_index: int, markdown_text: str) -> bool:
        """
        Replaces the content of a paragraph with new markdown text, applying formatting.
        """
        if target_index < 0 or target_index >= len(doc.paragraphs):
            return False
            
        para = doc.paragraphs[target_index]
        logging.info(f"apply_markdown: applying to paragraph {target_index}")
        apply_markdown_to_paragraph(para, markdown_text)
        return True

    def set_east_asian_font(self, run, font_name):
        """
        Safely sets the East Asian font for a run.
        """
        from docx.oxml.ns import qn
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


    def save_to_stream(self):
        target_doc = self.staging_doc if self.staging_doc else self.doc
        if not target_doc:
            return None
        stream = io.BytesIO()
        target_doc.save(stream)
        stream.seek(0)
        return stream
