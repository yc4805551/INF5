import re
from docx.text.paragraph import Paragraph

def apply_markdown_to_paragraph(paragraph: Paragraph, markdown_text: str):
    """
    Parses markdown text (bold **text**, italic *text*) and applies it to the paragraph.
    Clears existing runs and creates new ones.
    """
    # Clear existing content
    for run in paragraph.runs:
        run.text = ""
    
    # Simple parser for **bold** and *italic*
    # We'll use a regex to tokenize
    # Token types: BOLD, ITALIC, TEXT
    
    # Regex for **bold**
    # Regex for *italic*
    # We need to handle them carefully.
    # Let's use a simple split approach for now.
    
    # Pattern: (**.*?**)|(\*.*?\*)
    pattern = r"(\*\*.*?\*\*)|(\*.*?\*)"
    parts = re.split(pattern, markdown_text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith("**") and part.endswith("**"):
            # Bold
            text = part[2:-2]
            run = paragraph.add_run(text)
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            # Italic
            text = part[1:-1]
            run = paragraph.add_run(text)
            run.italic = True
        else:
            # Normal text
            paragraph.add_run(part)
