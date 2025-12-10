import os
import logging
import json
import re
import hashlib
import uuid
import time
import io
import requests
import httpx
import difflib
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Any
from flask import Flask, request, jsonify, Response, stream_with_context
from flask_cors import CORS
from dotenv import load_dotenv
from pymilvus import connections, Collection, utility, FieldSchema, DataType, CollectionSchema
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import click

# Import DocxEngine from core
from core.docx_engine import DocxEngine
from core.llm_engine import LLMEngine

dotenv_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'config', '.env')
load_dotenv(dotenv_path)

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Constants
MODEL_MAPPING = {
    "kb_qwen": "qwen3-embedding:0.6b",
    "kb_nomic": "nomic-embed-text"
}
DEFAULT_EMBEDDING_MODEL = "nomic-embed-text"

# Ali/Qwen Configuration
ALI_API_KEY = os.getenv("ALI_API_KEY")
ALI_TARGET_URL = os.getenv("ALI_TARGET_URL", "https://dashscope.aliyuncs.com/compatible-mode")
ALI_MODEL = os.getenv("ALI_MODEL", "qwen-plus")

MILVUS_HOST = os.getenv("MILVUS_HOST", "127.0.0.1")
# 自动清理 http/https 前缀，防止连接失败
if MILVUS_HOST.startswith("http://"):
    MILVUS_HOST = MILVUS_HOST.replace("http://", "")
elif MILVUS_HOST.startswith("https://"):
    MILVUS_HOST = MILVUS_HOST.replace("https://", "")

MILVUS_PORT_STR = os.getenv("MILVUS_PORT", "19530")
try:
    MILVUS_PORT = int(MILVUS_PORT_STR)
except ValueError:
    logging.warning(f"MILVUS_PORT '{MILVUS_PORT_STR}' 无效, 使用默认值 19530")
    MILVUS_PORT = 19530

# --- Ollama 配置修复 (解决 502 Bad Gateway) ---
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://127.0.0.1")

# 1. 强制修复 Windows 下无法访问 0.0.0.0 的问题
if "0.0.0.0" in OLLAMA_HOST:
    logging.warning(f"检测到 OLLAMA_HOST 为 0.0.0.0，正在自动转换为 127.0.0.1 以确保连接成功。")
    OLLAMA_HOST = OLLAMA_HOST.replace("0.0.0.0", "127.0.0.1")

# 2. 补全协议头
if not OLLAMA_HOST.startswith("http"):
    OLLAMA_HOST = f"http://{OLLAMA_HOST}"

OLLAMA_PORT = os.getenv("OLLAMA_PORT", "11434")
OLLAMA_EMBED_API_URL = f"{OLLAMA_HOST}:{OLLAMA_PORT}/api/embeddings"
logging.info(f"Ollama API 地址已配置为: {OLLAMA_EMBED_API_URL}")

# --- 知识库路径 ---
KNOWLEDGE_BASE_DIR = os.getenv("KNOWLEDGE_BASE_DIR", "./knowledge_base")
KNOWLEDGE_BASE_DIR_NOMIC = os.getenv("KNOWLEDGE_BASE_DIR_NOMIC", "./knowledge_base_nomic")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 500))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 50))
INGEST_WORKERS = int(os.getenv("INGEST_WORKERS", 8))

# --- AI 代理配置 ---

# Gemini (OpenAI 兼容模式)
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("VITE_GEMINI_API_KEY")
GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/openai"
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

# OpenAI
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "dummy-key-for-proxy")
OPENAI_TARGET_URL = os.getenv("OPENAI_TARGET_URL", "https://api.chatanywhere.tech")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5") # 注意: gpt-5 目前可能只是占位符或特定渠道模型
# DeepSeek
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "dummy-key-for-proxy")
DEEPSEEK_TARGET_URL = os.getenv("DEEPSEEK_TARGET_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
DEEPSEEK_ENDPOINT = os.getenv("DEEPSEEK_ENDPOINT", "https://api.deepseek.com/v1/chat/completions")

current_engine = DocxEngine()
llm_engine = LLMEngine()
def get_model_for_collection(collection_name: str) -> str:
    for key, model_name in MODEL_MAPPING.items():
        if key in collection_name:
            logging.info(f"集合 '{collection_name}' 匹配到关键词 '{key}'，使用模型: '{model_name}'")
            return model_name
    logging.info(f"集合 '{collection_name}' 未匹配到任何关键词，使用默认模型: '{DEFAULT_EMBEDDING_MODEL}'")
    return DEFAULT_EMBEDDING_MODEL

def get_ollama_embedding(text: str, model_name: str):
    try:
        payload = {"model": model_name, "prompt": text}
        response = requests.post(OLLAMA_EMBED_API_URL, json=payload, timeout=60)
        response.raise_for_status()
        response_data = response.json()
        if "embedding" not in response_data:
            raise ValueError(f"Ollama API 响应 (模型: {model_name}) 中缺少 'embedding' 字段。")
        return response_data["embedding"]
    except requests.exceptions.RequestException as e:
        print(f"DEBUG: 调用 Ollama API (模型: {model_name}) 失败: {e}")
        logging.error(f"调用 Ollama API (模型: {model_name}) 失败: {e}")
        raise RuntimeError(f"无法连接到 Ollama 服务: {e}")
    except Exception as e:
        print(f"DEBUG: 从 Ollama (模型: {model_name}) 获取嵌入时出错: {e}")
        logging.error(f"从 Ollama (模型: {model_name}) 获取嵌入时出错: {e}")
        raise

def create_milvus_collection(collection_name, dim):
    if utility.has_collection(collection_name):
        return Collection(collection_name)
    logging.info(f"集合 '{collection_name}' 不存在，正在创建 (维度: {dim})...")
    fields = [
        FieldSchema(name="id", dtype=DataType.VARCHAR, max_length=36, is_primary=True),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="source_file", dtype=DataType.VARCHAR, max_length=1024),
        FieldSchema(name="chunk_index", dtype=DataType.INT64),
        FieldSchema(name="full_path", dtype=DataType.VARCHAR, max_length=4096),
        FieldSchema(name="embedding", dtype=DataType.FLOAT_VECTOR, dim=dim)
    ]
    schema = CollectionSchema(fields, description=f"知识库集合: {collection_name}")
    collection = Collection(name=collection_name, schema=schema)
    index_params = {"metric_type": "COSINE", "index_type": "IVF_FLAT", "params": {"nlist": 128}}
    collection.create_index(field_name="embedding", index_params=index_params)
    return collection

def text_to_chunks(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    if not text: return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def get_file_hash(file_path):
    """计算文件内容的哈希值，用于检测文件是否真正发生变化"""
    sha256_hash = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    except Exception as e:
        logging.error(f"计算文件哈希值时出错: {e}")
        return None

def upsert_file_to_milvus(file_path: str, collection_name: str, model_name: str):
    filename = os.path.basename(file_path)
    print(f"DEBUG: 进入 upsert_file_to_milvus, 文件: {filename}")
    try:
        current_hash = get_file_hash(file_path)
        if not os.path.exists(file_path):
            return
        
        with open(file_path, 'r', encoding='utf-8') as f: 
            content = f.read()
        if not content.strip():
            return
            
        collection = Collection(collection_name)
        collection.load()
        
        # 简单检查内容是否重复 (基于已有逻辑)
        try:
            expr = f"source_file == '{filename}'"
            result = collection.query(expr, output_fields=["text"])
            if len(result) > 0 and current_hash:
                first_chunk = result[0].get("text", "")
                if content.startswith(first_chunk) and len(content) > 0:
                    logging.info(f"文件 '{filename}' 内容似乎未变化，跳过更新。")
                    return
        except Exception as e:
            logging.warning(f"检查文件存在性时出错: {e}")
            
        # 删除旧数据
        delete_expr = f"source_file == '{filename}'"
        collection.delete(delete_expr)
        
        chunks = text_to_chunks(content)
        if not chunks: return
        
        entities_to_insert = []
        print(f"DEBUG: 正在为 '{filename}' 处理 {len(chunks)} 个文本块...")
        logging.info(f"为 '{filename}' 处理 {len(chunks)} 个文本块...")
        with ThreadPoolExecutor(max_workers=INGEST_WORKERS) as executor:
            future_to_chunk = {executor.submit(get_ollama_embedding, chunk, model_name): (i, chunk) for i, chunk in enumerate(chunks)}
            for future in as_completed(future_to_chunk):
                try:
                    embedding = future.result()
                    i, chunk = future_to_chunk[future]
                    print(f"DEBUG: 成功获取 chunk {i} 的嵌入")
                    entity = {
                        "id": str(uuid.uuid4()),
                        "text": chunk,
                        "source_file": filename,
                        "chunk_index": i,
                        "full_path": file_path,
                        "embedding": embedding
                    }
                    entities_to_insert.append(entity)
                except Exception as e:
                    print(f"DEBUG: 生成嵌入失败: {e}")
                    logging.error(f"生成嵌入失败: {e}")
        if entities_to_insert:
            collection.insert(entities_to_insert)
            collection.flush()
            print(f"DEBUG: 已更新文件: {filename}")
            logging.info(f"已更新文件: {filename}")
    except Exception as e:
        print(f"DEBUG: 处理文件 '{filename}' 失败: {e}")
        logging.error(f"处理文件 '{filename}' 失败: {e}")
    finally:
        if 'collection' in locals():
            collection.release()

def process_file_delete(file_path, collection_name):
    filename = os.path.basename(file_path)
    logging.info(f"检测到文件删除: {filename}")
    try:
        collection = Collection(collection_name)
        collection.load()
        delete_expr = f"source_file == '{filename}'"
        collection.delete(delete_expr)
        logging.info(f"已删除索引: {filename}")
    except Exception as e:
        logging.error(f"删除索引失败: {e}")
    finally:
        if 'collection' in locals():
            collection.release()

# --- 文件监控处理器 ---
class KnowledgeBaseEventHandler(FileSystemEventHandler):
    def __init__(self, collection_to_watch, model_name, base_dir=None):
        self.collection_to_watch = collection_to_watch
        self.model_name = model_name
        self.base_dir = base_dir or KNOWLEDGE_BASE_DIR
        self.watch_path = os.path.normpath(os.path.join(self.base_dir, self.collection_to_watch))
        logging.info(f"监控器已初始化，目标路径: {self.watch_path}")
    def process_if_relevant(self, event):
        print(f"DEBUG: 收到事件: {event.event_type}, 路径: {event.src_path}")
        logging.info(f"收到事件: {event.event_type}, 路径: {event.src_path}")
        if event.is_directory:
            print("DEBUG: 忽略目录事件")
            logging.info("忽略目录事件")
            return
        if not (event.src_path.endswith(".txt") or event.src_path.endswith(".md")):
            print(f"DEBUG: 忽略非 txt/md 文件: {event.src_path}")
            logging.info("忽略非 txt/md 文件")
            return
            
        event_dir = os.path.normpath(os.path.dirname(event.src_path))
        # Case-insensitive comparison for Windows
        if event_dir.lower() != self.watch_path.lower():
            print(f"DEBUG: 路径不匹配: 事件目录 '{event_dir}' != 监控目录 '{self.watch_path}'")
            logging.info(f"路径不匹配: 事件目录 '{event_dir}' != 监控目录 '{self.watch_path}'")
            return
            
        print(f"DEBUG: 开始处理文件: {event.src_path}")
        logging.info(f"处理文件: {event.src_path}")
        if event.event_type in ('created', 'modified'):
            upsert_file_to_milvus(event.src_path, self.collection_to_watch, self.model_name)
        elif event.event_type == 'deleted':
            process_file_delete(event.src_path, self.collection_to_watch)
        elif event.event_type == 'moved':
            process_file_delete(event.src_path, self.collection_to_watch)
            dest_dir = os.path.normpath(os.path.dirname(event.dest_path))
            if dest_dir.lower() == self.watch_path.lower():
                upsert_file_to_milvus(event.dest_path, self.collection_to_watch, self.model_name)
    def on_created(self, event): self.process_if_relevant(event)
    def on_modified(self, event): self.process_if_relevant(event)
    def on_deleted(self, event): self.process_if_relevant(event)
    def on_moved(self, event): self.process_if_relevant(event)

# endregion

# region CLI Commands
def ingest_data():
    if not os.path.exists(KNOWLEDGE_BASE_DIR) or not os.path.isdir(KNOWLEDGE_BASE_DIR):
        logging.error(f"知识库根目录 '{KNOWLEDGE_BASE_DIR}' 不存在。")
        return
    for collection_name in os.listdir(KNOWLEDGE_BASE_DIR):
        collection_path = os.path.join(KNOWLEDGE_BASE_DIR, collection_name)
        if not os.path.isdir(collection_path): continue
        logging.info(f"\n--- 正在处理集合: {collection_name} ---")
        model_to_use = get_model_for_collection(collection_name)
        try:
            dummy_embedding = get_ollama_embedding("test", model_to_use)
            dim = len(dummy_embedding)
        except Exception as e:
            logging.error(f"无法获取模型维度，跳过。错误: {e}")
            continue
        create_milvus_collection(collection_name, dim)
        for filename in os.listdir(collection_path):
            file_path = os.path.join(collection_path, filename)
            if not (filename.endswith(".txt") or filename.endswith(".md")): continue
            upsert_file_to_milvus(file_path, collection_name, model_to_use)

@app.cli.command("ingest")
def ingest_command():
    ingest_data()
    click.echo("数据导入完成。")

@app.cli.command("watch")
def watch_command():
    collection_to_watch = 'kb_qwen_0_6b'
    connections.connect("default", host=MILVUS_HOST, port=MILVUS_PORT)
    model_name = get_model_for_collection(collection_to_watch)
    if not utility.has_collection(collection_to_watch):
        click.echo(f"错误: 集合 '{collection_to_watch}' 不存在。请先运行 'flask ingest'。")
        return
    path_to_watch = KNOWLEDGE_BASE_DIR
    event_handler = KnowledgeBaseEventHandler(collection_to_watch, model_name)
    observer = Observer()
    observer.schedule(event_handler, path_to_watch, recursive=True)
    click.echo(f"✅ 正在监控: {collection_to_watch}")
    observer.start()
    try:
        while True: time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

@app.cli.command("watch-nomic")
def watch_nomic_command():
    collection_to_watch = 'kb_nomic'
    connections.connect("default", host=MILVUS_HOST, port=MILVUS_PORT)
    model_name = get_model_for_collection(collection_to_watch)
    if not utility.has_collection(collection_to_watch):
        click.echo(f"错误: 集合 '{collection_to_watch}' 不存在。")
        return
    path_to_watch = KNOWLEDGE_BASE_DIR_NOMIC
    event_handler = KnowledgeBaseEventHandler(collection_to_watch, model_name, base_dir=KNOWLEDGE_BASE_DIR_NOMIC)
    observer = Observer()
    observer.schedule(event_handler, path_to_watch, recursive=True)
    click.echo(f"✅ 正在监控: {collection_to_watch}")
    observer.start()
    try:
        while True: time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()
# endregion

# region API Endpoints - Knowledge Base
@app.route('/api/list-collections', methods=['GET'])
def list_collections():
    try:
        try:
            collections = utility.list_collections()
        except Exception:
            logging.warning("Milvus 连接重试中...")
            try:
                connections.connect("default", host=MILVUS_HOST, port=MILVUS_PORT)
                collections = utility.list_collections()
            except Exception as e:
                logging.error(f"Milvus 连接失败: {e}")
                # Return empty list instead of crashing, so frontend can load
                return jsonify({"collections": [], "error": "Milvus unavailable"})
        return jsonify({"collections": collections})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/find-related', methods=['POST'])
def find_related():
    try:
        data = request.get_json()
        query_text = data.get('text')
        collection_name = data.get('collection_name')
        top_k = data.get('top_k', 10)
        
        if not query_text or not collection_name:
            return jsonify({"error": "缺少 text 或 collection_name"}), 400
        
        # Ensure connection exists before checking collection
        try:
            connections.connect("default", host=MILVUS_HOST, port=MILVUS_PORT)
        except Exception as e:
            logging.error(f"Milvus 连接失败: {e}")
            return jsonify({"error": f"无法连接到向量数据库: {str(e)}"}), 500

        if not utility.has_collection(collection_name):
            return jsonify({"error": f"集合 '{collection_name}' 不存在"}), 404

        model_to_use = get_model_for_collection(collection_name)
        query_embedding = get_ollama_embedding(query_text, model_to_use)
        
        collection = Collection(collection_name)
        collection.load()
        search_params = {"metric_type": "COSINE", "params": {"nprobe": 10}}
        results = collection.search(data=[query_embedding], anns_field="embedding", param=search_params, limit=top_k, 
                                    output_fields=["text", "source_file", "chunk_index", "full_path"])
        
        response_data = []
        for hit in results[0]:
            response_data.append({
                "source_file": hit.entity.get("source_file"),
                "content_chunk": hit.entity.get("text"),
                "score": hit.distance,
            })
        collection.release()
        return jsonify({"related_documents": response_data})
    except Exception as e:
        logging.error(f"API /find-related 错误: {e}")
        return jsonify({"error": str(e)}), 500
# endregion

# region API Endpoints - Canvas (Conversational Editor)
@app.route('/api/canvas/upload', methods=['POST'])
def canvas_upload():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        content = file.read()
        current_engine.load_document(io.BytesIO(content))
        return jsonify({"message": "File uploaded successfully", "preview": current_engine.get_preview_data()})
    except Exception as e:
        logging.error(f"Canvas Upload Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/preview', methods=['GET'])
def canvas_preview():
    try:
        return jsonify(current_engine.get_preview_data())
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/chat', methods=['POST'])
def canvas_chat():
    try:
        data = request.get_json()
        user_text = data.get("message")
        model_config = data.get("model_config", {})
        history = data.get("history", [])
        selection_context = data.get("selection_context", [])

        if not user_text:
            return jsonify({"error": "Message required"}), 400

        # Create staging copy for this new interaction if needed
        if not current_engine.staging_doc:
            current_engine.create_staging_copy()

        # Get current context from the STAGING doc
        context = current_engine.get_preview_data()

        # Interact with LLM
        response = llm_engine.chat_with_doc(
            user_message=user_text, 
            doc_context=context, 
            model_config=model_config,
            history=history,
            selection_context=selection_context
        )

        intent = response.get("intent")
        reply = response.get("reply") or ""
        code = response.get("code")
        
        is_staging = False
        
        if intent == "MODIFY":
            if code:
                # Execute Code on STAGING
                success, error_msg = current_engine.execute_code(code)
                
                if not success:
                     # If execution fails, return error in reply and downgrade to CHAT
                     reply = f"{reply}\n\n(Error executing changes: {error_msg})"
                     intent = "CHAT"
                else:
                     is_staging = True
            else:
                # MODIFY intent but no code? Downgrade to CHAT
                intent = "CHAT"
        
        return jsonify({
            "message": "Processed", 
            "reply": reply,
            "intent": intent,
            "preview": current_engine.get_preview_data(),
            "is_staging": is_staging
        })
    except Exception as e:
        logging.error(f"Canvas Chat Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/confirm', methods=['POST'])
def canvas_confirm():
    try:
        success = current_engine.commit_staging()
        if not success:
            return jsonify({"error": "No pending changes to confirm"}), 400
        return jsonify({
            "message": "Changes confirmed",
            "preview": current_engine.get_preview_data()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/discard', methods=['POST'])
def canvas_discard():
    try:
        current_engine.discard_staging()
        return jsonify({
            "message": "Changes discarded",
            "preview": current_engine.get_preview_data()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/download', methods=['GET'])
def canvas_download():
    try:
        stream = current_engine.save_to_stream()
        return Response(
            stream, 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=modified.docx"}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/format_official', methods=['POST'])
def canvas_format_official():
    try:
        data = request.get_json()
        model_config = data.get("model_config", {})
        scope = data.get("scope", "all")
        processor = data.get("processor", "local") # Default to local for speed

        # Ensure staging exists
        if not current_engine.staging_doc:
            current_engine.create_staging_copy()

        # Get context
        context = current_engine.get_preview_data()

        # Generate Formatting Code
        code = llm_engine.generate_formatting_code(context, model_config, scope, processor)
        
        # Execute Code on STAGING
        success, error_msg = current_engine.execute_code(code)
        
        if not success:
             return jsonify({"error": f"Failed to execute formatting code: {error_msg}"}), 400

        return jsonify({
            "message": "Formatting applied", 
            "code_executed": code,
            "preview": current_engine.get_preview_data(),
            "is_staging": True
        })
    except Exception as e:
        logging.error(f"Canvas Format Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/patch', methods=['POST'])
def canvas_patch():
    try:
        patch = request.get_json()
        # For direct patch, let's assume it applies to committed doc for now
        success = current_engine.apply_patch(patch, use_staging=False)
        if not success:
            return jsonify({"error": "Failed to apply patch"}), 400
        return jsonify({"message": "Patch applied", "preview": current_engine.get_preview_data()})
    except Exception as e:
        logging.error(f"Canvas Patch Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/canvas/modify_local', methods=['POST'])
def canvas_modify_local():
    try:
        data = request.get_json()
        file_path = data.get("file_path")
        instruction = data.get("instruction")
        
        if not file_path or not instruction:
            return jsonify({"error": "file_path and instruction are required"}), 400
        
        # 1. Load
        current_engine.load_from_path(file_path)
        
        # 2. Generate Code (using committed doc context)
        context = current_engine.get_preview_data()
        code = llm_engine.generate_code(instruction, context)
        
        # 3. Execute Code
        success, error_msg = current_engine.execute_code(code)
        if not success:
             return jsonify({"error": f"Failed to execute AI code: {error_msg}"}), 500
        
        # 4. Save
        current_engine.save_to_path(file_path)
        
        return jsonify({
            "message": "File processed and saved",
            "file_path": file_path,
            "preview": current_engine.get_preview_data()
        })
    except Exception as e:
        logging.error(f"Canvas Modify Local Error: {e}")
        return jsonify({"error": str(e)}), 500
# endregion

# region API Endpoints - AI Agents
@app.route('/api/generate', methods=['POST']) 
def handle_generate(): 
    try: 
        data = request.get_json() 
        provider = data.get('provider') 
        logging.info(f"收到非流式生成请求，Provider: {provider}") 
        
        if provider == 'gemini': 
            if not GEMINI_API_KEY: return jsonify({"error": "GEMINI_API_KEY 未设置"}), 500 
            return _call_gemini_openai_proxy(data)
        elif provider == 'openai': 
            return _call_openai_proxy(data) 
        elif provider == 'deepseek': 
            return _call_deepseek_proxy(data) 
        elif provider == 'ali': 
            return _call_ali_proxy(data) 
        else: 
            return jsonify({"error": f"不支持的 provider: {provider}"}), 400 
    except Exception as e: 
        logging.error(f"API /api/generate 错误: {e}") 
        return jsonify({"error": str(e)}), 500 

@app.route('/api/generate-stream', methods=['POST']) 
def handle_generate_stream(): 
    try: 
        data = request.get_json() 
        provider = data.get('provider') 
        sys_inst = data.get('systemInstruction') 
        user_prompt = data.get('userPrompt') 
        history = data.get('history', []) 
        
        logging.info(f"收到流式生成请求，Provider: {provider}") 

        if provider == 'gemini': 
            if not GEMINI_API_KEY: return Response(stream_with_context(["[错误: GEMINI_API_KEY 未设置]"]), content_type='text/plain')
            return Response(stream_with_context(_stream_gemini_openai_proxy(user_prompt, sys_inst, history)), content_type='text/plain') 
        elif provider == 'openai': 
            return Response(stream_with_context(_stream_openai_proxy(user_prompt, sys_inst, history)), content_type='text/plain') 
        elif provider == 'deepseek': 
            return Response(stream_with_context(_stream_deepseek_proxy(user_prompt, sys_inst, history)), content_type='text/plain') 
        elif provider == 'ali': 
            return Response(stream_with_context(_stream_ali_proxy(user_prompt, sys_inst, history)), content_type='text/plain') 
        else: 
            return Response(stream_with_context([f"[错误: 不支持的 provider: {provider}]"]), content_type='text/plain') 
    except Exception as e: 
        return Response(stream_with_context([f"[内部错误: {str(e)}]"]), content_type='text/plain')

# region Proxy Implementations

# --- 1. Gemini (OpenAI Compatible) ---
def _call_gemini_openai_proxy(data):
    base_url = GEMINI_BASE_URL.rstrip('/')
    # [修复] 移除多余的 /v1，因为 base_url 已经是 v1beta/openai
    url = f"{base_url}/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {GEMINI_API_KEY}'}
    
    messages = []
    if data.get('systemInstruction'):
        messages.append({"role": "system", "content": data.get('systemInstruction')})
    for item in data.get('history', []):
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": data.get('userPrompt')})
    
    payload = {"model": GEMINI_MODEL, "messages": messages, "temperature": 0.7}
    
    response = requests.post(url, headers=headers, json=payload, timeout=180)
    if response.status_code != 200:
        logging.error(f"Gemini 代理错误 ({response.status_code}): {response.text}")
        response.raise_for_status()
        
    response_data = response.json()
    if 'choices' in response_data and len(response_data['choices']) > 0:
        return Response(response_data['choices'][0]['message']['content'], content_type='application/json')
    raise ValueError("Gemini 响应格式无效")

def _stream_gemini_openai_proxy(user_prompt, system_instruction, history):
    base_url = GEMINI_BASE_URL.rstrip('/')
    url = f"{base_url}/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {GEMINI_API_KEY}'}
    
    messages = []
    if system_instruction: messages.append({"role": "system", "content": system_instruction})
    for item in history:
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": user_prompt})
    
    payload = {"model": GEMINI_MODEL, "messages": messages, "temperature": 0.7, "stream": True}
    
    try:
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=180) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if line:
                    line_str = line.decode('utf-8').strip()
                    if line_str.startswith('data: '):
                        line_str = line_str[6:]
                        if line_str == '[DONE]': break
                        try:
                            chunk = json.loads(line_str)
                            if chunk.get('choices') and chunk['choices'][0].get('delta', {}).get('content'):
                                yield chunk['choices'][0]['delta']['content']
                        except: pass
    except Exception as e:
        logging.error(f"Gemini 流式错误: {e}")
        yield f"[代理错误: {str(e)}]"

# --- 2. OpenAI ---
def _call_openai_proxy(data):
    url = f"{OPENAI_TARGET_URL}/v1/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {OPENAI_API_KEY}'}
    messages = []
    if data.get('systemInstruction'): messages.append({"role": "system", "content": data.get('systemInstruction')})
    for item in data.get('history', []):
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": data.get('userPrompt')})
    payload = {"model": OPENAI_MODEL, "messages": messages, "temperature": 0.7}
    response = requests.post(url, headers=headers, json=payload, timeout=180)
    response.raise_for_status()
    return Response(response.json()['choices'][0]['message']['content'], content_type='application/json')

def _stream_openai_proxy(user_prompt, system_instruction, history):
    url = f"{OPENAI_TARGET_URL}/v1/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {OPENAI_API_KEY}'}
    messages = []
    if system_instruction: messages.append({"role": "system", "content": system_instruction})
    for item in history:
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": user_prompt})
    payload = {"model": OPENAI_MODEL, "messages": messages, "temperature": 0.7, "stream": True}
    try:
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=180) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if line:
                    line_str = line.decode('utf-8').strip()
                    if line_str.startswith('data: '):
                        line_str = line_str[6:]
                        if line_str == '[DONE]': break
                        try:
                            chunk = json.loads(line_str)
                            if chunk.get('choices') and chunk['choices'][0].get('delta', {}).get('content'):
                                yield chunk['choices'][0]['delta']['content']
                        except: pass
    except Exception as e: yield f"[代理错误: {str(e)}]"

# --- 3. DeepSeek ---
def _call_deepseek_proxy(data):
    url = DEEPSEEK_ENDPOINT
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {DEEPSEEK_API_KEY}'}
    messages = []
    if data.get('systemInstruction'): messages.append({"role": "system", "content": data.get('systemInstruction')})
    for item in data.get('history', []):
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": data.get('userPrompt')})
    payload = {"model": DEEPSEEK_MODEL, "messages": messages, "temperature": 0.7}
    response = requests.post(url, headers=headers, json=payload, timeout=180)
    response.raise_for_status()
    return Response(response.json()['choices'][0]['message']['content'], content_type='application/json')

def _stream_deepseek_proxy(user_prompt, system_instruction, history):
    url = DEEPSEEK_ENDPOINT
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {DEEPSEEK_API_KEY}'}
    messages = []
    if system_instruction: messages.append({"role": "system", "content": system_instruction})
    for item in history:
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": user_prompt})
    payload = {"model": DEEPSEEK_MODEL, "messages": messages, "temperature": 0.7, "stream": True}
    try:
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=180) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if line:
                    line_str = line.decode('utf-8').strip()
                    if line_str.startswith('data: '):
                        line_str = line_str[6:]
                        if line_str == '[DONE]': break
                        try:
                            chunk = json.loads(line_str)
                            if chunk.get('choices') and chunk['choices'][0].get('delta', {}).get('content'):
                                yield chunk['choices'][0]['delta']['content']
                        except: pass
    except Exception as e: yield f"[代理错误: {str(e)}]"

# --- 4. Ali (Doubao/Qwen) ---
def _call_ali_proxy(data):
    url = f"{ALI_TARGET_URL}/v1/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {ALI_API_KEY}'}
    messages = []
    if data.get('systemInstruction'): messages.append({"role": "system", "content": data.get('systemInstruction')})
    for item in data.get('history', []):
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": data.get('userPrompt')})
    payload = {"model": ALI_MODEL, "messages": messages, "temperature": 0.7}
    response = requests.post(url, headers=headers, json=payload, timeout=180)
    response.raise_for_status()
    return Response(response.json()['choices'][0]['message']['content'], content_type='application/json')

def _stream_ali_proxy(user_prompt, system_instruction, history):
    url = f"{ALI_TARGET_URL}/v1/chat/completions"
    headers = {'Content-Type': 'application/json', 'Authorization': f'Bearer {ALI_API_KEY}'}
    messages = []
    if system_instruction: messages.append({"role": "system", "content": system_instruction})
    for item in history:
        if item.get('role') and item.get('parts'):
            messages.append({"role": item.get('role'), "content": item.get('parts')[0].get('text')})
    messages.append({"role": "user", "content": user_prompt})
    payload = {"model": ALI_MODEL, "messages": messages, "temperature": 0.7, "stream": True}
    try:
        with requests.post(url, headers=headers, json=payload, stream=True, timeout=180) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if line:
                    line_str = line.decode('utf-8').strip()
                    if line_str.startswith('data: '):
                        line_str = line_str[6:]
                        if line_str == '[DONE]': break
                        try:
                            chunk = json.loads(line_str)
                            if chunk.get('choices') and chunk['choices'][0].get('delta', {}).get('content'):
                                yield chunk['choices'][0]['delta']['content']
                        except: pass
    except Exception as e: yield f"[代理错误: {str(e)}]"

# endregion

@app.route('/api/canvas/export_docx', methods=['POST'])
def export_docs():
    data = request.json
    markdown_content = data.get('markdown', '')

    try:
        doc = Document()
        
        # Setup Page Margins
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        def set_font(run, font_name, size_pt, bold=False):
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.bold = bold

        lines = markdown_content.split('\n')
        first_line_processed = False

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue

            clean_text = stripped_line.replace('*', '').replace('#', '').strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28) 

            if not first_line_processed:
                # TITLE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(clean_text)
                set_font(run, '小标宋体', 22, bold=False)
                first_line_processed = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                is_l1 = re.match(r'^[一二三四五六七八九十]+、', clean_text)
                is_l2 = re.match(r'^（[一二三四五六七八九十]+）', clean_text)
                
                if is_l1:
                    run = p.add_run(clean_text)
                    set_font(run, '黑体', 16)
                elif is_l2:
                    run = p.add_run(clean_text)
                    set_font(run, '楷体_GB2312', 16)
                else:
                    p.paragraph_format.first_line_indent = Cm(1.1) 
                    run = p.add_run(clean_text)
                    set_font(run, '仿宋_GB2312', 16)

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return Response(
            f.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-disposition": "attachment; filename=export.docx"}
        )

    except Exception as e:
        logging.error(f"Generate DOCX failed: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5179)
